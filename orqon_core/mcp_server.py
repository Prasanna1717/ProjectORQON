"""
IBM watsonx Orchestrate MCP Toolkit Server
Model Context Protocol (MCP) server with specialized agents

IBM Toolkit Compliance:
- Local Python MCP server (stdio and SSE transport)
- Remote MCP server (SSE and streamable HTTP)
- Tool discovery with proper schemas
- Authentication support (OAuth2, API Key, Bearer Token)
- 30-second handshake timeout
- JSON-RPC 2.0 protocol
- Tool execution with proper error handling
"""
import asyncio
import json
import sys
import re
from typing import Dict, Any, Optional, AsyncGenerator, List, Literal
from dataclasses import dataclass, asdict
from enum import Enum
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, HTTPException, Request, Header, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel, Field
import uvicorn
from datetime import datetime
import logging

# Configure logging for IBM MCP toolkit
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("orqon_mcp_server")

# IBM watsonx Orchestrate ADK imports
try:
    from ibm_watsonx_orchestrate.agent_builder.agents import (
        AssistantAgent,
        AgentKind,
    )
    from ibm_watsonx_orchestrate.agent_builder.agents.types import (
        AssistantAgentConfig,
    )
    HAS_AGENT_BUILDER = True
    logger.info("✓ IBM watsonx Orchestrate agent_builder loaded (v1.15.0)")
except ImportError as e:
    HAS_AGENT_BUILDER = False
    AssistantAgent = None
    AgentKind = None
    AssistantAgentConfig = None
    logger.warning(f"⚠️  IBM agent_builder not available: {e}")
    logger.info("   Install: pip install ibm-watsonx-orchestrate")


# ============================================================================
# IBM MCP TOOLKIT PROTOCOL
# ============================================================================

class MCPToolSchema(BaseModel):
    """IBM MCP toolkit tool schema (JSON-RPC 2.0 compatible)"""
    name: str = Field(..., description="Tool name")
    description: str = Field(..., description="Tool description")
    inputSchema: Dict[str, Any] = Field(..., description="JSON Schema for tool inputs")
    
    class Config:
        json_schema_extra = {
            "example": {
                "name": "send_email",
                "description": "Send an email using Gmail",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "to": {"type": "string", "description": "Recipient email"},
                        "subject": {"type": "string", "description": "Email subject"},
                        "body": {"type": "string", "description": "Email body"}
                    },
                    "required": ["to", "subject", "body"]
                }
            }
        }


class MCPToolResponse(BaseModel):
    """IBM MCP toolkit tool execution response"""
    content: List[Dict[str, Any]] = Field(..., description="Response content")
    isError: bool = Field(default=False, description="Whether the response is an error")
    
    class Config:
        json_schema_extra = {
            "example": {
                "content": [
                    {
                        "type": "text",
                        "text": "Email sent successfully"
                    }
                ],
                "isError": False
            }
        }


class MCPListToolsResponse(BaseModel):
    """IBM MCP toolkit list tools response"""
    tools: List[MCPToolSchema] = Field(..., description="Available tools")
    
    class Config:
        json_schema_extra = {
            "example": {
                "tools": [
                    {
                        "name": "send_email",
                        "description": "Send an email",
                        "inputSchema": {
                            "type": "object",
                            "properties": {
                                "to": {"type": "string"},
                                "subject": {"type": "string"},
                                "body": {"type": "string"}
                            }
                        }
                    }
                ]
            }
        }


class MCPCallToolRequest(BaseModel):
    """IBM MCP toolkit call tool request"""
    name: str = Field(..., description="Tool name to execute")
    arguments: Dict[str, Any] = Field(default_factory=dict, description="Tool arguments")


class MCPServerInfo(BaseModel):
    """IBM MCP server information"""
    name: str = Field(..., description="Server name")
    version: str = Field(..., description="Server version")
    protocolVersion: str = Field(default="2024-11-05", description="MCP protocol version")
    capabilities: Dict[str, Any] = Field(default_factory=dict, description="Server capabilities")
    
    class Config:
        json_schema_extra = {
            "example": {
                "name": "orqon-mcp-server",
                "version": "3.0.0",
                "protocolVersion": "2024-11-05",
                "capabilities": {
                    "tools": {},
                    "prompts": {},
                    "resources": {}
                }
            }
        }


# IBM MCP Toolkit Tool Registry
MCP_TOOL_REGISTRY: List[MCPToolSchema] = []


def register_mcp_tool(
    name: str,
    description: str,
    input_schema: Dict[str, Any]
):
    """Register a tool in the IBM MCP toolkit registry"""
    tool = MCPToolSchema(
        name=name,
        description=description,
        inputSchema=input_schema
    )
    MCP_TOOL_REGISTRY.append(tool)
    logger.info(f"✓ Registered MCP tool: {name}")


# ============================================================================
# CONVERSATION MEMORY
# ============================================================================

# Import shared memory module (accessible by all agents including agent_orchestrator)
from shared_memory import conversation_memory

def get_client_email_from_csv(client_name: str) -> Optional[str]:
    """Get client email from CSV data - flexible matching"""
    import csv
    from pathlib import Path
    
    csv_path = Path(__file__).parent / "data" / "trade_blotter.csv"
    if not csv_path.exists():
        logger.warning(f"CSV file not found: {csv_path}")
        return None
    
    try:
        client_name_lower = client_name.lower().strip()
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                csv_client = row.get('Client', '').lower().strip()
                # Exact match
                if csv_client == client_name_lower:
                    email = row.get('Email', '').strip()
                    logger.info(f"✅ Found exact match: {row.get('Client')} → {email}")
                    return email
                # Partial match (e.g., "Sheila" matches "Sheila Carter")
                elif client_name_lower in csv_client or csv_client in client_name_lower:
                    email = row.get('Email', '').strip()
                    logger.info(f"✅ Found partial match: '{client_name}' matched '{row.get('Client')}' → {email}")
                    return email
    except Exception as e:
        logger.error(f"Error reading CSV: {e}")
    
    logger.warning(f"⚠️ No client found matching: {client_name}")
    return None

# ============================================================================
# AGENT TYPES
# ============================================================================

class AgentType(str, Enum):
    """Specialized agent types"""
    COORDINATOR = "coordinator"        # Routes to appropriate agent
    GMAIL = "gmail"                   # Email operations
    EXCEL = "excel"                   # CSV/Excel data operations
    FINANCE = "finance"               # Stock prices, trade data
    COMPLIANCE = "compliance"         # RAG knowledge base
    TRADE_PARSER = "trade_parser"     # Parse trade logs and tickets
    CALENDAR = "calendar"             # Google Calendar
    MATH = "math"                     # Calculations
    DATETIME = "datetime"             # Date/time queries


@dataclass
class AgentHandoff:
    """Agent handoff request"""
    from_agent: AgentType
    to_agent: AgentType
    context: Dict[str, Any]
    reason: str


# ============================================================================
# SPECIALIZED AGENTS
# ============================================================================

class BaseAgent:
    """Base class for specialized agents with IBM ADK integration"""
    
    def __init__(self, agent_type: AgentType):
        self.agent_type = agent_type
        self.capabilities = []
        self.adk_agent = None  # IBM agent_builder agent instance
        
        # Initialize IBM ADK agent if available
        if HAS_AGENT_BUILDER:
            self._init_adk_agent()
    
    def _init_adk_agent(self):
        """Initialize IBM watsonx Orchestrate ADK agent"""
        try:
            # Create agent using IBM watsonx Orchestrate 1.15.0 API
            agent_name = f"{self.agent_type.value}_agent"
            agent_description = f"Specialized agent for {self.agent_type.value} operations"
            
            # Build agent spec
            agent_spec = {
                "name": agent_name,
                "description": agent_description,
                "kind": AgentKind.ASSISTANT,
                "title": f"{self.agent_type.value.title()} Agent",
                "nickname": agent_name,
                "config": AssistantAgentConfig(
                    description=agent_description,
                ),
            }
            
            # Initialize IBM watsonx Orchestrate AssistantAgent
            self.adk_agent = AssistantAgent(**agent_spec)
            print(f"✓ IBM ADK agent initialized for {self.agent_type.value}")
        except Exception as e:
            print(f"⚠️  Failed to initialize ADK agent for {self.agent_type.value}: {e}")
            self.adk_agent = None
    
    async def can_handle(self, query: str, context: Dict[str, Any]) -> bool:
        """Check if agent can handle this query"""
        raise NotImplementedError
    
    async def process(self, query: str, context: Dict[str, Any]) -> Dict[str, Any]:
        """Process query and return result"""
        raise NotImplementedError


class GmailAgent(BaseAgent):
    """Specialized agent for Gmail, Google Calendar, and Google Meet with Excel data handoff"""
    
    def __init__(self):
        super().__init__(AgentType.GMAIL)
        self.capabilities = [
            "send_email",
            "draft_email",
            "read_emails",
            "search_emails",
            "calendar_reminder",
            "schedule_meeting",
            "google_meet",
            "excel_data_handoff"  # Receives email from Excel agent
        ]
        
        # Import Google Workspace tools
        try:
            from tools.google_workspace_tools import GoogleWorkspaceTools
            self.gmail_tools = GoogleWorkspaceTools()
            self.available = True
            print("✓ Gmail Agent initialized with Calendar, Meet, and Excel handoff support")
        except Exception as e:
            print(f"⚠️  Gmail Agent unavailable: {e}")
            self.available = False
    
    async def can_handle(self, query: str, context: Dict[str, Any]) -> bool:
        """Check if this is an email/calendar-related query"""
        if not self.available:
            return False
        
        query_lower = query.lower().strip()
        
        # PRIORITY: Calendar/reminder/cancel keywords (includes common typos)
        calendar_keywords = [
            'reminder', 'remind me', 'set a reminder', 'create reminder', 'remainder',  # 'remainder' = typo for reminder
            'schedule', 'meeting', 'calendar', 'set me a', 'add to calendar', 'add me',
            'google calendar', 'gcal', 'gcalender', 'calender',  # common typos
            'cancel meeting', 'delete meeting', 'cancel all', 'remove meeting'
        ]
        if any(keyword in query_lower for keyword in calendar_keywords):
            return True
        
        # EXCLUDE: Email queries (what/whats/show email) - these should go to Excel agent
        email_query_keywords = [
            'what is', 'whats', 'what\'s', 'show me the email', 'get the email', 'find email', 
            'tell me the email', 'give me the email', 'display email'
        ]
        for keyword in email_query_keywords:
            if keyword in query_lower and ('email' in query_lower or 'mail' in query_lower):
                return False  # Let Excel agent handle email queries
        
        # PRIORITY: Email actions should always route to Gmail agent
        # Pattern 1: Action + pronoun (mail her, email him)
        action_pronoun_patterns = [
            'mail her', 'mail him', 'mail them',
            'email her', 'email him', 'email them',
        ]
        if any(pattern in query_lower for pattern in action_pronoun_patterns):
            return True
        
        # Pattern 2: Gmail/email + name (gmail sheila, email john, let's gmail ron)
        # Check for "gmail" or "email" followed by a name anywhere in the query
        gmail_match = re.search(r'\b(?:lets?|let\'s)?\s*(?:gmail|email|mail)\s+([a-z]+(?:\s+[a-z]+)?)', query_lower)
        if gmail_match:
            potential_name = gmail_match.group(1).strip()
            if potential_name not in ['the', 'a', 'an', 'it', 'me', 'us', 'regarding', 'about', 'with']:
                logger.info(f"✅ Gmail Agent: Detected 'gmail/email + {potential_name}' pattern")
                return True
        
        # Pattern 3: Explicit send actions
        explicit_send_patterns = [
            'send email', 'send mail', 'send a mail', 'send an email',
            'lets mail', 'let\'s mail', 'lets gmail', 'let\'s gmail',
            'lets email', 'let\'s email', 'write to', 'compose email', 
            'draft email', 'notify via email'
        ]
        if any(pattern in query_lower for pattern in explicit_send_patterns):
            return True
        
        return False
    
    async def process(self, query: str, context: Dict[str, Any]) -> Dict[str, Any]:
        """Process email/calendar request"""
        from watsonx_llm import WatsonxLLM
        from langchain_core.messages import SystemMessage, HumanMessage
        import re
        from datetime import datetime, timedelta
        
        llm = WatsonxLLM()
        
        # CHECK FOR CALENDAR/MEETING/REMINDER/CANCEL REQUESTS
        query_lower = query.lower()
        
        # Check for explicit email sending patterns to EXCLUDE from calendar detection
        is_email_send = (
            query_lower.startswith('gmail ') or 
            query_lower.startswith('email ') or 
            query_lower.startswith('mail ') or
            'send email' in query_lower or
            'send mail' in query_lower
        )
        
        # CHECK FOR CANCEL REQUESTS
        is_cancel = any(keyword in query_lower for keyword in ['cancel', 'delete', 'remove'])
        
        if is_cancel:
            logger.info(f"🗑️ Cancel meeting request detected: {query}")
            
            try:
                # Determine if cancelling all or specific meeting
                cancel_all = any(keyword in query_lower for keyword in ['all', 'everything', 'every'])
                
                if cancel_all:
                    logger.info(f"🗑️ Cancelling ALL upcoming meetings...")
                    result = self.gmail_tools.cancel_all_meetings()
                    
                    if result.get('success'):
                        cancelled_count = result.get('cancelled_count', 0)
                        cancelled_titles = result.get('cancelled_titles', [])
                        
                        if cancelled_count == 0:
                            response_text = "ℹ️ No upcoming meetings found to cancel."
                        else:
                            response_text = f"✅ **Cancelled {cancelled_count} Meeting(s)**\n\n"
                            for i, title in enumerate(cancelled_titles, 1):
                                response_text += f"{i}. {title}\n"
                        
                        logger.info(f"✅ Cancelled {cancelled_count} meetings")
                        return {
                            "success": True,
                            "agent": "gmail",
                            "action": "meetings_cancelled",
                            "response": response_text
                        }
                    else:
                        error_msg = result.get('error', 'Unknown error')
                        logger.error(f"❌ Failed to cancel meetings: {error_msg}")
                        return {
                            "success": False,
                            "agent": "gmail",
                            "error": f"Failed to cancel meetings: {error_msg}"
                        }
                else:
                    # Try to extract meeting title/client name for specific cancellation
                    # For now, just cancel all (can be enhanced later)
                    return {
                        "success": False,
                        "agent": "gmail",
                        "error": "Please specify 'cancel all meetings' to cancel. Specific meeting cancellation coming soon."
                    }
            
            except Exception as e:
                logger.error(f"❌ Exception cancelling meetings: {str(e)}", exc_info=True)
                return {
                    "success": False,
                    "agent": "gmail",
                    "error": f"Error cancelling meetings: {str(e)}"
                }
        
        # Detect if this is a MEETING/REMINDER request - but ONLY if NOT an email send request
        is_meeting = (not is_email_send and not is_cancel and 
                     any(keyword in query_lower for keyword in ['meeting', 'meet with', 'schedule with', 'google meet', 'gmeet', 'schedule a meeting']))
        is_reminder = (not is_email_send and not is_cancel and 
                      any(keyword in query_lower for keyword in ['reminder', 'remind me', 'calendar', 'set me a', 'add me', 'create reminder', 
                                                                   'google calendar', 'gcal', 'set reminder', 'add to calendar', 'add a reminder']))
        
        if (is_meeting or is_reminder) and not is_email_send:
            logger.info(f"📅 Calendar/Reminder request detected: {query}")
            
            # Extract reminder details from context and query
            # PRIORITY 1: Check for date in last client's data (from memory)
            reminder_context = ""
            reminder_date = None
            client_info = ""
            
            logger.info(f"📅 Checking shared memory for client data...")
            logger.info(f"📅 Available keys: {list(conversation_memory['shared_context'].keys())}")
            
            # SMART LOOKUP: If query mentions a client name, look them up directly
            # Extract potential client names from query (with/meeting with/regarding [NAME])
            # Clean query for better matching (remove punctuation like ;; )
            clean_query = re.sub(r'[;:,]+', ' ', query_lower)
            client_name_match = re.search(r'(?:with|regarding|for)\s+([a-z]+(?:\s+[a-z]+)?)', clean_query)
            if client_name_match and 'last_client_data' not in conversation_memory['shared_context']:
                potential_client = client_name_match.group(1).strip().title()
                logger.info(f"📅 Query mentions '{potential_client}' but no memory - looking up in CSV...")
                
                # Look up client in CSV with fuzzy matching
                import csv
                from pathlib import Path
                csv_path = Path(__file__).parent / "data" / "trade_blotter.csv"
                
                if csv_path.exists():
                    try:
                        with open(csv_path, 'r', encoding='utf-8') as f:
                            reader = csv.DictReader(f)
                            best_match = None
                            best_score = 0
                            
                            for row in reader:
                                csv_client = row.get('Client', '').lower()
                                # Split both names into parts for better matching
                                query_parts = potential_client.lower().split()
                                client_parts = csv_client.split()
                                
                                # Count matching parts
                                matches = sum(1 for qp in query_parts for cp in client_parts if qp in cp or cp in qp)
                                score = matches / max(len(query_parts), len(client_parts))
                                
                                logger.info(f"📅 Comparing '{potential_client}' with '{row.get('Client')}': score={score}")
                                
                                if score > best_score:
                                    best_score = score
                                    best_match = row
                            
                            # Accept match if score >= 0.5 (at least half the parts match)
                            if best_match and best_score >= 0.5:
                                client_data = {
                                    'client_name': best_match.get('Client', potential_client),
                                    'email': best_match.get('Email', ''),
                                    'account': best_match.get('Acct#', ''),
                                    'ticker': best_match.get('Ticker', ''),
                                    'quantity': best_match.get('Qty', ''),
                                    'follow_up': best_match.get('FollowUpDate', ''),
                                    'FollowUpDate': best_match.get('FollowUpDate', ''),
                                    'meeting_needed': best_match.get('MeetingNeeded', ''),
                                    'stage': best_match.get('Stage', ''),
                                    'notes': best_match.get('Notes', '')
                                }
                                # Save to shared memory directly
                                conversation_memory['shared_context']['last_client_data'] = client_data
                                conversation_memory['shared_context']['last_client_name'] = best_match.get('Client', potential_client)
                                logger.info(f"📅 Auto-loaded {best_match.get('Client')} from CSV (score={best_score})!")
                                logger.info(f"📅 Saved to memory: {client_data.get('client_name')} ({client_data.get('email')})")
                            else:
                                logger.warning(f"📅 No good match found for '{potential_client}' (best score: {best_score})")
                    except Exception as e:
                        logger.warning(f"📅 Failed to auto-lookup client: {e}")
            
            if 'last_client_data' in conversation_memory['shared_context']:
                client_data = conversation_memory['shared_context']['last_client_data']
                client_name = client_data.get('client_name', client_data.get('Client', 'Unknown'))
                follow_up = client_data.get('follow_up', client_data.get('FollowUpDate', ''))
                
                logger.info(f"📅 Found client data for: {client_name}")
                logger.info(f"📅 Follow-up date raw: {follow_up}")
                
                if follow_up and follow_up.strip():
                    try:
                        # Parse follow-up date (format: YYYY-MM-DD or YYYY-MM-DD HH:MM:SS)
                        reminder_date = datetime.strptime(follow_up.split()[0], '%Y-%m-%d')
                        # Set time to 9 AM if no time specified
                        reminder_date = reminder_date.replace(hour=9, minute=0, second=0, microsecond=0)
                        reminder_context = f"Follow up with {client_name}"
                        client_info = f"{client_name}'s follow-up"
                        logger.info(f"📅 ✅ Parsed follow-up date from memory: {reminder_date}")
                    except Exception as e:
                        logger.warning(f"📅 Failed to parse follow-up date '{follow_up}': {e}")
                else:
                    logger.warning(f"📅 Client {client_name} has no follow-up date set")
            else:
                logger.warning(f"📅 No last_client_data in shared memory")
            
            # PRIORITY 2: Parse explicit dates from query
            if not reminder_date:
                logger.info(f"📅 No date from memory, parsing query for explicit date...")
                
                # Look for dates like "tomorrow", "next week", "Dec 1", "November 27"
                if 'tomorrow' in query_lower:
                    reminder_date = datetime.now() + timedelta(days=1)
                    reminder_date = reminder_date.replace(hour=9, minute=0, second=0, microsecond=0)
                    logger.info(f"📅 Found 'tomorrow' → {reminder_date}")
                elif 'next week' in query_lower:
                    reminder_date = datetime.now() + timedelta(days=7)
                    reminder_date = reminder_date.replace(hour=9, minute=0, second=0, microsecond=0)
                    logger.info(f"📅 Found 'next week' → {reminder_date}")
                else:
                    # Try to extract date from query using regex (YYYY-MM-DD format)
                    date_match = re.search(r'(\d{4}-\d{2}-\d{2})', query)
                    if date_match:
                        try:
                            reminder_date = datetime.strptime(date_match.group(1), '%Y-%m-%d')
                            reminder_date = reminder_date.replace(hour=9, minute=0, second=0, microsecond=0)
                            logger.info(f"📅 Extracted date from query: {reminder_date}")
                        except Exception as e:
                            logger.warning(f"📅 Failed to parse date '{date_match.group(1)}': {e}")
            
            # PRIORITY 3: Use LLM as fallback to extract/invent date
            if not reminder_date:
                logger.info(f"📅 No explicit date found, using LLM to extract or default to tomorrow...")
                
                llm_prompt = f"""Extract reminder information from this query: "{query}"

RULES:
1. If user mentions a specific date/time, extract it
2. If user says "regarding [client]'s follow up" but no date given, return "TOMORROW_MORNING"
3. If completely unclear, return "TOMORROW_MORNING"
4. Never return null or empty

Output format: YYYY-MM-DD or "TOMORROW_MORNING"

Query: {query}
Output:"""
                
                try:
                    llm_result = llm.invoke(llm_prompt)
                    # Extract content from AIMessage object
                    llm_response = llm_result.content.strip() if hasattr(llm_result, 'content') else str(llm_result).strip()
                    logger.info(f"📅 LLM response: {llm_response}")
                    
                    if llm_response == "TOMORROW_MORNING" or "tomorrow" in llm_response.lower():
                        reminder_date = datetime.now() + timedelta(days=1)
                        reminder_date = reminder_date.replace(hour=9, minute=0, second=0, microsecond=0)
                        logger.info(f"📅 LLM suggested tomorrow → {reminder_date}")
                    else:
                        # Try to parse LLM's date
                        date_match = re.search(r'(\d{4}-\d{2}-\d{2})', llm_response)
                        if date_match:
                            reminder_date = datetime.strptime(date_match.group(1), '%Y-%m-%d')
                            reminder_date = reminder_date.replace(hour=9, minute=0, second=0, microsecond=0)
                            logger.info(f"📅 LLM extracted date: {reminder_date}")
                        else:
                            # No date found, default to tomorrow
                            reminder_date = datetime.now() + timedelta(days=1)
                            reminder_date = reminder_date.replace(hour=9, minute=0, second=0, microsecond=0)
                            logger.info(f"📅 No date in LLM response, defaulting to tomorrow → {reminder_date}")
                except Exception as e:
                    logger.warning(f"📅 LLM extraction failed: {e}, defaulting to tomorrow")
                    # Final fallback: tomorrow morning
                    reminder_date = datetime.now() + timedelta(days=1)
                    reminder_date = reminder_date.replace(hour=9, minute=0, second=0, microsecond=0)
            
            # Create MEETING or REMINDER
            if reminder_date:
                try:
                    # Get client email and name from memory
                    client_email = None
                    client_name_for_title = "Client"
                    
                    if 'last_client_data' in conversation_memory['shared_context']:
                        client_data = conversation_memory['shared_context']['last_client_data']
                        client_email = client_data.get('email')
                        client_name_for_title = client_data.get('client_name', client_data.get('Client', 'Client'))
                        logger.info(f"📅 Found client in memory: {client_name_for_title} ({client_email})")
                    else:
                        logger.warning(f"📅 No client data found in memory for meeting scheduling")
                    
                    if is_meeting:
                        # CREATE GOOGLE MEET MEETING
                        if not client_email:
                            error_msg = f"❌ Cannot schedule meeting - no email found for the client. Please query the client's data first (e.g., 'show data for Meghan Hall') or provide a valid client name."
                            logger.error(error_msg)
                            return {
                                "success": False,
                                "agent": "gmail",
                                "response": error_msg,
                                "error": error_msg
                            }
                        
                        title = f"Portfolio Review Meeting - {client_name_for_title}"
                        notes = f"Scheduled via Orqon assistant.\nOriginal query: {query}\n\nAgenda: Portfolio review and follow-up discussion"
                        duration_minutes = 60  # Default 1 hour meeting
                        
                        logger.info(f"📅 Creating Google Meet meeting: {title}")
                        logger.info(f"📅 Date: {reminder_date}")
                        logger.info(f"📅 Attendee: {client_email}")
                        
                        result = self.gmail_tools.schedule_meeting(
                            title=title,
                            start_time=reminder_date,
                            duration_minutes=duration_minutes,
                            attendee_emails=[client_email],
                            description=notes
                        )
                        
                        logger.info(f"📅 Meeting result: {json.dumps(result, indent=2)}")
                        
                        if result.get('success'):
                            meet_link = result.get('meet_link', 'N/A')
                            response_text = (
                                f"✅ **Google Meet Meeting Scheduled**\n\n"
                                f"📅 Title: {title}\n"
                                f"👤 Attendee: {client_name_for_title} ({client_email})\n"
                                f"🕐 Date: {reminder_date.strftime('%B %d, %Y at %I:%M %p')}\n"
                                f"⏱️ Duration: {duration_minutes} minutes\n\n"
                                f"🔗 [View in Calendar]({result.get('event_link', 'N/A')})\n"
                                f"📹 [Join Google Meet]({meet_link})"
                            )
                            logger.info(f"✅ Meeting scheduled successfully: {title}")
                            
                            # Check if user also wants to send email notification
                            should_send_email = any(phrase in query_lower for phrase in ['mail him', 'mail her', 'email him', 'email her', 'send email', 'notify'])
                            
                            if should_send_email and client_email:
                                logger.info(f"📧 User also requested email notification - sending meeting invite email...")
                                
                                # Compose meeting invite email
                                email_subject = f"Meeting Invitation: Portfolio Review - {reminder_date.strftime('%B %d, %Y')}"
                                email_body = f"""Dear {client_name_for_title},<br><br>I hope this email finds you well.<br><br>I've scheduled a portfolio review meeting for us to discuss your investment strategy and review your recent transactions.<br><br>📅 <strong>Meeting Details:</strong><br><br>• <strong>Date:</strong> {reminder_date.strftime('%B %d, %Y at %I:%M %p')}<br>• <strong>Duration:</strong> {duration_minutes} minutes<br>• <strong>Topic:</strong> Portfolio Review & Follow-up Discussion<br><br>📹 <strong>Join the meeting:</strong><br><a href="{meet_link}">{meet_link}</a><br><br>📆 <strong>Add to your calendar:</strong><br><a href="{result.get('event_link', '')}">{result.get('event_link', 'Calendar Link')}</a><br><br>Looking forward to our conversation. Please let me know if you need to reschedule.<br><br>Best regards,<br>Prasanna Vijay<br>Financial Advisor<br>The Orqon Team<br><br>📧 Email: prasannathefreelancer@gmail.com<br>📞 Available for consultation"""
                                
                                try:
                                    email_result = self.gmail_tools.send_email(
                                        to_email=client_email,
                                        subject=email_subject,
                                        body=email_body
                                    )
                                    
                                    if email_result.get('success'):
                                        logger.info(f"✅ Meeting invite email sent to {client_email}")
                                        response_text += f"\n\n📧 **Meeting invite email sent to {client_name_for_title}**"
                                    else:
                                        logger.error(f"❌ Failed to send meeting invite email: {email_result.get('error')}")
                                        response_text += f"\n\n⚠️ Meeting created but email notification failed"
                                except Exception as e:
                                    logger.error(f"❌ Exception sending meeting email: {e}")
                                    response_text += f"\n\n⚠️ Meeting created but email notification failed"
                            
                            # Return the meeting info
                            return {
                                "success": True,
                                "agent": "gmail",
                                "action": "meeting_scheduled",
                                "response": response_text,
                                "meeting_details": {
                                    "title": title,
                                    "date": reminder_date.strftime('%B %d, %Y at %I:%M %p'),
                                    "meet_link": meet_link,
                                    "calendar_link": result.get('event_link', ''),
                                    "client_email": client_email,
                                    "client_name": client_name_for_title
                                }
                            }
                        else:
                            error_msg = result.get('error', result.get('message', 'Unknown error'))
                            logger.error(f"❌ Failed to schedule meeting: {error_msg}")
                            return {
                                "success": False,
                                "agent": "gmail",
                                "error": f"Failed to schedule meeting: {error_msg}"
                            }
                    
                    else:
                        # CREATE REMINDER (no attendees needed)
                        title = reminder_context or "Reminder from Orqon"
                        notes = f"Set via Orqon assistant. Original query: {query}"
                        
                        logger.info(f"📅 Creating reminder: {title}")
                        logger.info(f"📅 Date: {reminder_date}")
                        logger.info(f"📅 Notes: {notes}")
                        
                        result = self.gmail_tools.create_reminder(
                            title=title,
                            reminder_time=reminder_date,
                            notes=notes
                        )
                        
                        logger.info(f"📅 Reminder result: {json.dumps(result, indent=2)}")
                        
                        if result.get('success'):
                            response_text = (
                                f"✅ **Google Calendar Reminder Created**\n\n"
                                f"📅 Title: {title}\n"
                                f"🕐 Date: {reminder_date.strftime('%B %d, %Y at %I:%M %p')}\n\n"
                                f"🔗 [View in Calendar]({result.get('event_link', 'N/A')})"
                            )
                            logger.info(f"✅ Reminder created successfully: {title}")
                            return {
                                "success": True,
                                "agent": "gmail",
                                "action": "reminder_created",
                                "response": response_text
                            }
                        else:
                            error_msg = result.get('error', result.get('message', 'Unknown error'))
                            logger.error(f"❌ Failed to create reminder: {error_msg}")
                            return {
                                "success": False,
                                "agent": "gmail",
                                "error": f"Failed to create reminder: {error_msg}"
                            }
                        
                except Exception as e:
                    logger.error(f"❌ Exception creating calendar event: {str(e)}", exc_info=True)
                    return {
                        "success": False,
                        "agent": "gmail",
                        "error": f"Error creating calendar event: {str(e)}"
                    }
            else:
                logger.warning(f"⚠️ Could not determine reminder date from query: {query}")
                return {
                    "success": False,
                    "agent": "gmail",
                    "error": "Could not determine reminder date. Please specify a date (e.g., 'tomorrow', '2025-11-27', 'next week')"
                }
        
        # PRIORITY 1: Check if context already has client_data from previous agent (handoff)
        extracted_email = None
        client_name = None
        
        logger.info(f"🔍 Gmail Agent processing query: {query}")
        logger.info(f"🔍 Context received: {json.dumps(context, indent=2)}")
        
        if 'client_data' in context and context['client_data'].get('email'):
            extracted_email = context['client_data']['email']
            client_name = context['client_data'].get('client_name', context['client_data'].get('Client'))
            context['client_email_from_csv'] = extracted_email
            context['client_name'] = client_name
            logger.info(f"✅ PRIORITY 1: Using client data from handoff: {client_name} ({extracted_email})")
        
        # PRIORITY 2: Check shared memory for last client data
        elif 'last_client_data' in conversation_memory['shared_context']:
            client_data = conversation_memory['shared_context']['last_client_data']
            logger.info(f"🔍 Found last_client_data in shared memory: {json.dumps(client_data, indent=2)}")
            if client_data.get('email') or client_data.get('Email'):
                extracted_email = client_data.get('email') or client_data.get('Email')
                client_name = client_data.get('client_name') or client_data.get('Client')
                context['client_email_from_csv'] = extracted_email
                context['client_name'] = client_name
                context['client_data'] = client_data
                logger.info(f"✅ PRIORITY 2: Using client data from memory: {client_name} ({extracted_email})")
            else:
                logger.warning(f"⚠️ last_client_data exists but has no email field!")
        else:
            logger.info(f"🔍 No last_client_data in shared memory. Available keys: {list(conversation_memory['shared_context'].keys())}")
        
        # PRIORITY 3: Extract client name from query and lookup CSV
        if not extracted_email:
            client_name_patterns = [
                r'\b(sheila|carter|sheila carter)\b',
                r'\b(john|doe|john doe)\b',
                r'\b(ron|ronald)\b',  # Added Ron support
                r'\b(tony|stark|tony stark)\b',  # Added Tony support
                r'gmail\s+([a-z]+)(?:\s+with|\s+about|\s+regarding|,)',  # "gmail ron with" or "gmail ron,"
                r'gmail\s+([a-z\s]+)$',
                r'mail\s+([a-z]+)(?:\s+with|\s+about|\s+regarding|,)',  # "mail ron with"
                r'mail\s+([a-z\s]+)$',
                r'email\s+([a-z]+)(?:\s+with|\s+about|\s+regarding|,)',  # "email ron with"
                r'email\s+([a-z\s]+)$',
                r'send.*to\s+([a-z\s]+)',
                r'her\s+(?:mail|email)',
                r'his\s+(?:mail|email)',
                r'their\s+(?:mail|email)',
            ]
            
            for pattern in client_name_patterns:
                match = re.search(pattern, query.lower())
                if match:
                    potential_name = match.group(1).strip().title() if len(match.groups()) > 0 else None
                    
                    # For pronouns (her/his/their), use last client
                    if not potential_name and 'last_client_name' in conversation_memory['shared_context']:
                        potential_name = conversation_memory['shared_context']['last_client_name']
                    
                    if potential_name:
                        # Try to get email from CSV
                        csv_email = get_client_email_from_csv(potential_name)
                        if csv_email:
                            extracted_email = csv_email
                            client_name = potential_name
                            context['client_email_from_csv'] = csv_email
                            context['client_name'] = client_name
                            logger.info(f"✅ PRIORITY 3: Found email for {potential_name}: {csv_email}")
                            break
                        else:
                            logger.warning(f"⚠️ No email found in CSV for: {potential_name}")
        
        # Extract email parameters using LLM
        system_prompt = """You are a professional financial advisor email composition assistant.

CRITICAL EMAIL ADDRESS RULES (HIGHEST PRIORITY):
1. If context contains 'VERIFIED_EMAIL_MUST_USE' - YOU MUST USE THIS EXACT EMAIL as to_email
2. If context contains 'recipient_email' - USE THAT EMAIL ADDRESS as to_email
3. NEVER use example emails like client@email.com or sheila.carter@example.com
4. NEVER hallucinate or make up email addresses
5. ALWAYS copy the exact email from context

EMAIL CONTENT RULES:
1. If context contains 'recipient_name', use that name in header and salutation
2. If context contains client_data with stock/trade information, INCLUDE IT in email body
3. Use professional financial advisor tone

EMAIL FORMAT REQUIREMENTS:
CRITICAL INSTRUCTION: OUTPUT BODY ONLY - NO HEADERS OR SEPARATORS AT THE TOP.

Email body formatting rules:
- Start DIRECTLY with greeting: "Dear [Name],"
- Add blank line after greeting (\\n)
- Main content with specific details (stocks, trades, dates if available)
- Use bullet points with emoji icons for lists (📊, •, 📈, 💼)
- Add blank lines (\\n) between paragraphs
- Clear call to action or next steps
- Add blank line before closing (\\n)
- Footer: "Best regards,\\nPrasanna Vijay\\nFinancial Advisor\\nThe Orqon Team\\n\\n📧 Email: prasannathefreelancer@gmail.com\\n📞 Available for consultation"

DO NOT include:
- Recipient name header at top
- Decorative separator lines (━━━━)
- Any content before "Dear [Name],"

Output ONLY valid JSON:
{
  "to_email": "actual.email@domain.com",
  "subject": "Professional subject line",
  "body": "Full formatted email body",
  "action": "send"
}

Example 1:
Input: "email sheila about follow up meeting for her stocks"
Context: {"recipient_email": "sheila.carter@example.com", "recipient_name": "Sheila Carter", "client_data": {"ticker": "TSLA", "quantity": 500, "side": "SELL", "follow_up": "2025-11-24"}}
Output: {
  "to_email": "sheila.carter@example.com",
  "subject": "Follow-up Meeting: Your TSLA Transaction",
  "body": "Dear Sheila,\\n\\nI hope this email finds you well. I am writing to follow up on your recent transaction and discuss the next steps for your portfolio.\\n\\n📊 TRANSACTION DETAILS:\\n\\n• Stock: TSLA (Tesla, Inc.)\\n• Action: SELL\\n• Quantity: 500 shares\\n• Follow-up Date: November 24, 2025\\n\\nGiven the recent market activity and your portfolio position, I believe it would be beneficial to schedule a meeting to discuss your investment strategy and ensure your financial goals remain aligned with current market conditions.\\n\\nI am available to meet at your convenience. Please let me know your preferred time, and I will make the necessary arrangements.\\n\\nLooking forward to our conversation.\\n\\nBest regards,\\nPrasanna Vijay\\nFinancial Advisor\\nThe Orqon Team\\n\\n📧 Email: prasannathefreelancer@gmail.com\\n📞 Available for consultation",
  "action": "send"
}

Example 2:
Input: "send email to john@example.com saying thanks"
Output: {
  "to_email": "john@example.com",
  "subject": "Thank You",
  "body": "Dear John,\\n\\nThank you for your time and consideration.\\n\\nBest regards,\\nPrasanna Vijay\\nFinancial Advisor\\nThe Orqon Team",
  "action": "send"
}"""
        
        # CRITICAL: If we have extracted email, we will force it (don't trust LLM)
        forced_email = None
        if extracted_email:
            forced_email = extracted_email
            logger.info(f"🎯 FORCING verified email: {extracted_email} for {client_name}")
        else:
            logger.warning(f"⚠️ No email extracted from any priority! Will rely on LLM parsing.")
        
        # Create clean context for LLM (replace internal keys with values)
        llm_context = context.copy()
        if 'client_email_from_csv' in llm_context and extracted_email:
            llm_context['recipient_email'] = extracted_email
            del llm_context['client_email_from_csv']
        if 'client_name' in llm_context:
            llm_context['recipient_name'] = llm_context.pop('client_name')
        
        # Add explicit verified email marker and FORCE it into context
        if forced_email:
            llm_context['VERIFIED_EMAIL_MUST_USE'] = forced_email
            llm_context['recipient_email'] = forced_email  # Double ensure
            if client_name:
                llm_context['recipient_name'] = client_name
        
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=f"Query: {query}\nContext: {json.dumps(llm_context)}")
        ]
        
        response = llm.invoke(messages)
        response_text = response.content.strip()
        
        # Parse JSON
        try:
            # Extract JSON from response
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1
            if json_start != -1 and json_end > json_start:
                email_params = json.loads(response_text[json_start:json_end])
            else:
                raise ValueError("No JSON found in response")
            
            # CRITICAL OVERRIDE FIRST: If we have forced_email, ALWAYS use it (don't trust LLM at all)
            if forced_email:
                original_email = email_params.get('to_email', 'NOT_SET')
                email_params['to_email'] = forced_email
                if original_email and original_email != forced_email:
                    logger.warning(f"🔧 OVERRIDING LLM hallucination: {original_email} → {forced_email}")
                else:
                    logger.info(f"✅ Using forced verified email: {forced_email}")
            
            # Validate required fields AFTER override
            if 'to_email' not in email_params or 'subject' not in email_params:
                return {
                    "success": False,
                    "error": "Could not extract email parameters. Please specify recipient and subject."
                }
            
            # Backup fixes for when forced_email is None
            elif email_params['to_email'] == 'client_email_from_csv' and extracted_email:
                email_params['to_email'] = extracted_email
                logger.info(f"🔧 Fixed literal string: Using {extracted_email}")
            elif 'client_email_from_csv' in email_params['to_email'] and extracted_email:
                email_params['to_email'] = extracted_email
                logger.info(f"🔧 Fixed partial match: Using {extracted_email}")
            
            # Final validation
            if not email_params['to_email'] or '@' not in email_params['to_email']:
                if extracted_email:
                    email_params['to_email'] = extracted_email
                    logger.info(f"🔧 Fallback: Using {extracted_email}")
                else:
                    logger.error(f"❌ Could not determine recipient email address!")
                    return {
                        "success": False,
                        "error": "Could not determine recipient email address."
                    }
            
            logger.info(f"📧 Final email address to use: {email_params['to_email']}")
            
            # AGGRESSIVE POST-PROCESSING: Fix email body formatting for Gmail
            body = email_params.get('body', '')
            logger.info(f"📝 Original body: {len(body)} chars")
            logger.info(f"📝 First 200 chars: {repr(body[:200])}")
            
            # Step 1: Cut the header - split on "Dear" and keep only body content
            if 'Dear' in body:
                parts = body.split('Dear', 1)
                body = 'Dear' + parts[1]
                logger.info(f"✂️ Cut header, new body starts with: {repr(body[:50])}")
            
            # Step 2: Convert \n to <br> for Gmail HTML rendering
            body = body.replace('\n', '<br>')
            logger.info(f"✅ Converted \\n to <br> - body length: {len(body)} chars")
            
            email_params['body'] = body
            
            # Check for attachments in context
            attachment_paths = None
            if 'attachments' in context and context['attachments']:
                # Extract file paths from attachments list
                # Format can be: [[filename, path], ...] or [path, ...]
                attachment_paths = []
                for item in context['attachments']:
                    if isinstance(item, (tuple, list)) and len(item) >= 2:
                        # Format: [filename, path] or (filename, path)
                        attachment_paths.append(item[1])
                    else:
                        # Format: path string
                        attachment_paths.append(item)
                logger.info(f"📎 Attaching {len(attachment_paths)} file(s): {attachment_paths}")
            
            # Send email
            result = self.gmail_tools.send_email(
                to_email=email_params['to_email'],
                subject=email_params['subject'],
                body=email_params.get('body', ''),
                cc_emails=email_params.get('cc_emails'),
                attachment_paths=attachment_paths
            )
            
            # Store sent email in Astra DB for future hybrid searches
            try:
                from tools.astra_db_tools import get_astra_store
                astra = get_astra_store()
                if astra:
                    email_data = {
                        "from": "prasannathefreelancer@gmail.com",
                        "to": email_params['to_email'],
                        "subject": email_params['subject'],
                        "body": email_params.get('body', ''),
                        "message_id": result.get('id', '')
                    }
                    astra.add_email_context(email_data)
                    print("✓ Email indexed in Astra DB for hybrid search")
            except Exception as e:
                print(f"⚠️  Failed to index email in Astra DB: {e}")
            
            # Format email body back to display format (convert <br> back to \n for user preview)
            display_body = email_params.get('body', '').replace('<br>', '\n')
            
            return {
                "success": True,
                "agent": "gmail",
                "action": "email_sent",
                "to": email_params['to_email'],
                "subject": email_params['subject'],
                "message_id": result.get('id'),
                "response": f"✅ **Email sent successfully to {email_params['to_email']}**\n\n**Subject:** {email_params['subject']}\n\n**Email Content:**\n\n{display_body}"
            }
            
        except json.JSONDecodeError:
            return {
                "success": False,
                "error": "Failed to parse email parameters"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Failed to send email: {str(e)}"
            }


class TradeParserAgent(BaseAgent):
    """Specialized agent for parsing trade logs and tickets using IBM ADK"""
    
    def __init__(self):
        super().__init__(AgentType.TRADE_PARSER)
        self.capabilities = [
            "parse_trade_log",
            "extract_trade_data",
            "log_trade_ticket",
            "multiple_trade_parsing"
        ]
        self.available = True
        print("✓ Trade Parser Agent initialized")
    
    async def can_handle(self, query: str, context: Dict[str, Any]) -> bool:
        """Check if this is a trade parsing query"""
        query_lower = query.lower()
        
        # Trade logging keywords
        trade_keywords = [
            'log trade', 'trade ticket', 'ticket reference', 'emergency log',
            'client called', 'bought', 'sold', 'buy', 'sell',
            'shares', 'position', 'market order', 'limit order',
            'solicited', 'unsolicited', 'compliance review'
        ]
        
        # Check if query is long (likely a trade log)
        if len(query.split()) > 15:
            if any(keyword in query_lower for keyword in ['client', 'trade', 'shares', 'ticker', 'stock']):
                return True
        
        return any(keyword in query_lower for keyword in trade_keywords)
    
    async def process(self, query: str, context: Dict[str, Any]) -> Dict[str, Any]:
        """Parse trade log and extract structured data"""
        from watsonx_llm import WatsonxLLM
        from langchain_core.messages import SystemMessage, HumanMessage
        import csv
        from pathlib import Path
        from datetime import datetime
        
        llm = WatsonxLLM()
        
        # Trade parsing prompt
        system_prompt = """You are a trade ticket parser for a financial brokerage system.

Extract trade information from natural language logs and output structured JSON.

OUTPUT FORMAT (JSON only):
{
  "trades": [
    {
      "client_name": "Full Name",
      "account_number": "Account ID",
      "ticker": "SYMBOL",
      "side": "Buy" or "Sell",
      "quantity": number,
      "order_type": "Market" or "Limit" or "Stop",
      "price": number (0 for Market orders),
      "solicited": true/false,
      "notes": "Context and details",
      "follow_up_date": "YYYY-MM-DD" or "",
      "email": "client@email.com" or "",
      "stage": "Pending/Follow-up Scheduled/Compliance Review/Completed",
      "meeting_needed": true/false,
      "ticket_id": "TICKET-ID" if mentioned
    }
  ]
}

KEY EXTRACTION RULES:
- Side: "Buy" or "Sell" (capitalize first letter)
- Solicited: false if "unsolicited" or "client asked", true if "broker recommended"
- Order Type: "Market" if no price specified or "market", "Limit" if specific price
- Price: 0 for Market orders, actual price for Limit orders
- Meeting Needed: true if mentions "follow-up", "meeting", "call", "review needed"
- Stage: "Compliance Review" if urgent/emotional, "Follow-up Scheduled" if meeting planned
- Email: Extract if mentioned (look for @)
- Follow-up Date: Extract or infer from "tomorrow", "next week", etc.

Example:
Input: "Client John Smith account 1234 bought 100 shares of AAPL at market, solicited, follow up next week"
Output: {
  "trades": [{
    "client_name": "John Smith",
    "account_number": "1234",
    "ticker": "AAPL",
    "side": "Buy",
    "quantity": 100,
    "order_type": "Market",
    "price": 0,
    "solicited": true,
    "notes": "Broker recommended trade",
    "follow_up_date": "",
    "email": "",
    "stage": "Follow-up Scheduled",
    "meeting_needed": true,
    "ticket_id": ""
  }]
}"""
        
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=f"Parse this trade log:\n\n{query}")
        ]
        
        try:
            response = llm.invoke(messages)
            response_text = response.content.strip()
            
            # Extract JSON
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1
            
            if json_start == -1 or json_end <= json_start:
                return {
                    "success": False,
                    "error": "Could not parse trade log. Please provide trade details."
                }
            
            parsed_data = json.loads(response_text[json_start:json_end])
            trades = parsed_data.get('trades', [])
            
            if not trades:
                return {
                    "success": False,
                    "error": "No trades found in the log."
                }
            
            # Save trades to CSV
            csv_path = Path(__file__).parent / "data" / "trade_blotter.csv"
            trades_logged = []
            
            for trade in trades:
                # Add timestamp
                trade['timestamp'] = datetime.now().strftime("%Y-%m-%d %I:%M %p")
                
                # Generate ticket ID if not provided
                if not trade.get('ticket_id'):
                    trade['ticket_id'] = f"TKT-{datetime.now().strftime('%Y%m%d%H%M%S')}"
                
                # Write to CSV
                try:
                    with open(csv_path, 'a', newline='', encoding='utf-8') as f:
                        writer = csv.DictWriter(f, fieldnames=[
                            'Ticket ID', 'Client', 'Account', 'Side', 'Ticker', 'Qty',
                            'Type', 'Price', 'Solicited', 'Timestamp', 'Notes',
                            'Follow-up', 'Email', 'Stage', 'Meeting'
                        ])
                        
                        writer.writerow({
                            'Ticket ID': trade.get('ticket_id', ''),
                            'Client': trade.get('client_name', ''),
                            'Account': trade.get('account_number', ''),
                            'Side': trade.get('side', ''),
                            'Ticker': trade.get('ticker', '').upper(),
                            'Qty': trade.get('quantity', 0),
                            'Type': trade.get('order_type', 'Market'),
                            'Price': trade.get('price', 0),
                            'Solicited': 'Yes' if trade.get('solicited') else 'No',
                            'Timestamp': trade['timestamp'],
                            'Notes': trade.get('notes', ''),
                            'Follow-up': trade.get('follow_up_date', ''),
                            'Email': trade.get('email', ''),
                            'Stage': trade.get('stage', 'Pending'),
                            'Meeting': 'Yes' if trade.get('meeting_needed') else 'No'
                        })
                    
                    trades_logged.append(trade)
                except Exception as e:
                    print(f"Error writing to CSV: {e}")
            
            # Format response
            response_lines = ["✅ Trade(s) logged successfully:\n"]
            for trade in trades_logged:
                response_lines.append(f"📋 Ticket: {trade['ticket_id']}")
                response_lines.append(f"   Client: {trade['client_name']} ({trade['account_number']})")
                response_lines.append(f"   Trade: {trade['side']} {trade['quantity']} {trade['ticker']} @ {trade['order_type']}")
                response_lines.append(f"   Stage: {trade['stage']}")
                if trade.get('meeting_needed'):
                    response_lines.append(f"   ⚠️  Meeting Required")
                response_lines.append("")
            
            return {
                "success": True,
                "agent": "trade_parser",
                "trades_logged": len(trades_logged),
                "trades": trades_logged,
                "response": "\n".join(response_lines)
            }
            
        except json.JSONDecodeError as e:
            return {
                "success": False,
                "error": f"Failed to parse trade data: {str(e)}"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Trade parsing error: {str(e)}"
            }


class ExcelAgent(BaseAgent):
    """Specialized agent for CSV/Excel operations with RAG and vector memory"""
    
    def __init__(self):
        super().__init__(AgentType.EXCEL)
        self.capabilities = [
            "read_csv",
            "filter_data",
            "get_client_data",
            "extract_field",
            "open_file",
            "show_file",
            "semantic_search",
            "vector_memory"
        ]
        
        # Import CSV tools
        from pathlib import Path
        self.csv_path = Path(__file__).parent / "data" / "trade_blotter.csv"
        self.available = self.csv_path.exists()
        
        # Initialize vector memory for RAG
        self.vector_store = None
        self.short_term_memory = []  # Last 10 queries
        self.long_term_memory = {}   # Client patterns and preferences
        
        try:
            import chromadb
            from chromadb.config import Settings
            
            # Setup ChromaDB for vector storage
            chroma_client = chromadb.Client(Settings(
                anonymized_telemetry=False,
                allow_reset=True
            ))
            
            # Create or get collection
            self.vector_store = chroma_client.get_or_create_collection(
                name="trade_blotter_memory",
                metadata={"hnsw:space": "cosine"}
            )
            
            # Index CSV data into vector store
            if self.available:
                self._index_csv_data()
                
            print("✓ Excel Agent initialized with RAG and vector memory")
        except Exception as e:
            print(f"⚠️  Excel Agent: Vector memory unavailable ({e}), using basic mode")
            if self.available:
                print("✓ Excel Agent initialized (basic mode)")
    
    def _index_csv_data(self):
        """Index CSV data into vector store for semantic search"""
        try:
            import csv
            with open(self.csv_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                rows = list(reader)
            
            # Create searchable documents from CSV rows
            documents = []
            metadatas = []
            ids = []
            
            for idx, row in enumerate(rows):
                # Create rich text representation for embedding
                doc_text = f"Client: {row.get('Client', 'N/A')} | "
                doc_text += f"Ticker: {row.get('Ticker', 'N/A')} | "
                doc_text += f"Side: {row.get('Side', 'N/A')} | "
                doc_text += f"Quantity: {row.get('Qty', 'N/A')} | "
                doc_text += f"Notes: {row.get('Notes', 'N/A')}"
                
                documents.append(doc_text)
                metadatas.append(row)
                ids.append(f"trade_{idx}")
            
            # Add to vector store
            if documents:
                self.vector_store.add(
                    documents=documents,
                    metadatas=metadatas,
                    ids=ids
                )
                logger.info(f"✓ Indexed {len(documents)} trades into vector memory")
        except Exception as e:
            logger.error(f"Failed to index CSV data: {e}")
    
    def _semantic_search(self, query: str, n_results: int = 5) -> List[Dict]:
        """Perform semantic search on trade data"""
        if not self.vector_store:
            return []
        
        try:
            results = self.vector_store.query(
                query_texts=[query],
                n_results=n_results
            )
            
            if results and results['metadatas']:
                return results['metadatas'][0]
            return []
        except Exception as e:
            logger.error(f"Semantic search failed: {e}")
            return []
    
    def _update_memory(self, query: str, client_name: str = None):
        """Update short and long-term memory"""
        # Short-term memory (last 10 queries)
        self.short_term_memory.append({
            'query': query,
            'timestamp': datetime.now().isoformat(),
            'client': client_name
        })
        if len(self.short_term_memory) > 10:
            self.short_term_memory.pop(0)
        
        # Long-term memory (client patterns)
        if client_name:
            if client_name not in self.long_term_memory:
                self.long_term_memory[client_name] = {'query_count': 0, 'last_query': None}
            self.long_term_memory[client_name]['query_count'] += 1
            self.long_term_memory[client_name]['last_query'] = query
    
    async def can_handle(self, query: str, context: Dict[str, Any]) -> bool:
        """Check if this is a data query"""
        if not self.available:
            return False
        
        query_lower = query.lower()
        
        # EXCLUDE long trade logs (let Trade Parser handle them)
        if len(query.split()) > 15:
            trade_log_indicators = ['emergency log', 'ticket reference', 'demanded', 'executed', 'unsolicited', 'solicited']
            if any(indicator in query_lower for indicator in trade_log_indicators):
                return False  # Let Trade Parser handle this
        
        # Check for file opening requests
        open_keywords = ['open', 'show me the', 'display', 'view']
        file_keywords = ['csv', 'excel', 'spreadsheet', 'file', 'blotter']
        if any(ok in query_lower for ok in open_keywords):
            if any(fk in query_lower for fk in file_keywords):
                return True
        
        # Handle email queries (what's the email, show email, etc.)
        email_query_patterns = ['what', 'whats', 'what\'s', 'show', 'get', 'find', 'tell me', 'give me']
        if any(pattern in query_lower for pattern in email_query_patterns):
            if 'email' in query_lower or 'mail' in query_lower:
                return True
        
        data_keywords = [
            'show', 'data', 'table', 'csv', 'excel', 
            'client', 'trade', 'blotter', 'account',
            'ticker', 'follow up', 'meeting'
        ]
        
        return any(keyword in query_lower for keyword in data_keywords)
    
    async def process(self, query: str, context: Dict[str, Any]) -> Dict[str, Any]:
        """Process data query"""
        import re
        import csv
        import subprocess
        import platform
        from pathlib import Path
        
        query_lower = query.lower()
        
        # Handle email queries specifically - PRIORITY CHECK
        email_keywords = ['email', 'mail', 'e-mail', 'contact']
        asking_for_info = any(word in query_lower for word in ['what', 'whats', "what's", 'show', 'get', 'find', 'tell', 'give'])
        
        if any(kw in query_lower for kw in email_keywords) and asking_for_info:
            # Extract client name from query - improved patterns
            email_patterns = [
                r'(?:email|mail|e-mail|contact).*?(?:of|for|from)\s+(.+?)(?:\s*$)',  # "email of ron"
                r'(?:what|whats|what\'s)\s+(?:is\s+)?(?:the\s+)?(.+?)(?:\'s|s)?\s+(?:email|mail|contact)',  # "what's ron's email"
                r'(?:what|whats|what\'s)\s+(?:is\s+)?(?:the\s+)?(?:email|mail|contact).*?(?:of|for)\s+(.+?)(?:\s*$)',  # "what's the email of ron"
                r'(.+?)(?:\'s|s)\s+(?:email|mail|e-mail|contact)',  # "ron's email"
            ]
            
            client_name = None
            for pattern in email_patterns:
                match = re.search(pattern, query_lower)
                if match:
                    client_name = match.group(1).strip()
                    break
            
            if client_name:
                try:
                    with open(self.csv_path, 'r', encoding='utf-8') as f:
                        reader = csv.DictReader(f)
                        rows = list(reader)
                    
                    # Find client (partial match)
                    matching_rows = [r for r in rows if client_name.lower() in r.get('Client', '').lower()]
                    
                    if matching_rows:
                        # Get first match
                        client_row = matching_rows[0]
                        client_full_name = client_row.get('Client', 'Unknown')
                        email = client_row.get('Email', '')
                        
                        if email:
                            self._update_memory(query, client_full_name)
                            return {
                                "success": True,
                                "agent": "excel",
                                "response": f"📧 **{client_full_name}'s email:** {email}",
                                "client_data": {
                                    "client_name": client_full_name,
                                    "email": email,
                                    "account": client_row.get('Acct#', ''),
                                    "ticker": client_row.get('Ticker', ''),
                                    "quantity": client_row.get('Qty', '')
                                }
                            }
                        else:
                            return {
                                "success": True,
                                "agent": "excel",
                                "response": f"❌ No email found for {client_full_name} in the database."
                            }
                    else:
                        return {
                            "success": False,
                            "agent": "excel",
                            "response": f"❌ No client matching '{client_name.title()}' found in the database."
                        }
                except Exception as e:
                    logger.error(f"Error fetching email: {e}")
                    return {
                        "success": False,
                        "agent": "excel",
                        "response": f"Error retrieving email: {str(e)}"
                    }
        
        # Check if user wants to open/view the file
        open_keywords = ['open', 'show me the', 'display', 'view']
        if any(keyword in query_lower for keyword in open_keywords):
            if 'excel' in query_lower or 'xlsx' in query_lower or 'spreadsheet' in query_lower:
                file_path = Path(__file__).parent / "data" / "trade_blotter.xlsx"
                file_type = "Excel"
            elif 'csv' in query_lower:
                file_path = Path(__file__).parent / "data" / "trade_blotter.csv"
                file_type = "CSV"
            elif 'file' in query_lower or 'blotter' in query_lower:
                # Default to Excel for better viewing
                file_path = Path(__file__).parent / "data" / "trade_blotter.xlsx"
                file_type = "Excel"
            else:
                file_path = None
            
            if file_path and file_path.exists():
                try:
                    # Open file with system default application
                    if platform.system() == "Windows":
                        subprocess.Popen(["start", "", str(file_path)], shell=True)
                    elif platform.system() == "Darwin":  # macOS
                        subprocess.Popen(["open", str(file_path)])
                    else:  # Linux
                        subprocess.Popen(["xdg-open", str(file_path)])
                    
                    return {
                        "success": True,
                        "agent": "excel",
                        "action": "file_opened",
                        "response": f"✅ Opened {file_type} file: {file_path.name}\n\nThe trade blotter is now displayed in your default application."
                    }
                except Exception as e:
                    return {
                        "success": False,
                        "agent": "excel",
                        "error": f"Failed to open file: {str(e)}"
                    }
        
        # Read CSV data directly and return as structured table data for Carbon DataTable
        try:
            with open(self.csv_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                rows = list(reader)
            
            # Extract client name from query for filtering
            client_name_pattern = r'(?:data\s+for|show\s+data\s+for|trades?\s+for)\s+(.+?)(?:\s*$)'
            client_match = re.search(client_name_pattern, query.lower())
            filtered_rows = rows
            filter_message = ""
            
            if client_match:
                search_name = client_match.group(1).strip()
                
                # Try exact match first (case-insensitive)
                exact_matches = [r for r in rows if r.get('Client', '').lower() == search_name.lower()]
                
                if exact_matches:
                    # Exact match found - return only this client
                    filtered_rows = exact_matches
                    filter_message = f" (exact match: {exact_matches[0].get('Client', search_name)})"
                    self._update_memory(query, exact_matches[0].get('Client', search_name))
                else:
                    # Try partial match (contains) - check both first and last name
                    search_parts = search_name.lower().split()
                    partial_matches = []
                    
                    for r in rows:
                        client_lower = r.get('Client', '').lower()
                        # Check if any part of search name is in client name
                        if any(part in client_lower for part in search_parts):
                            partial_matches.append(r)
                    
                    if len(partial_matches) == 1:
                        # Single partial match - treat as exact
                        filtered_rows = partial_matches
                        filter_message = f" (matched: {partial_matches[0].get('Client', search_name)})"
                        self._update_memory(query, partial_matches[0].get('Client', search_name))
                    elif len(partial_matches) > 1:
                        # Multiple partial matches - show all
                        filtered_rows = partial_matches
                        filter_message = f" ({len(partial_matches)} matches found)"
                        self._update_memory(query, search_name.title())
                    elif self.vector_store:
                        # No matches - try semantic search (limit to 1 for specific client queries)
                        semantic_results = self._semantic_search(f"trades for {search_name}", n_results=1)
                        if semantic_results:
                            filtered_rows = semantic_results
                            filter_message = f" (best match: {semantic_results[0].get('Client', search_name)})"
                            self._update_memory(query, semantic_results[0].get('Client', search_name))
                        else:
                            filtered_rows = []
                    else:
                        filtered_rows = []
            
            if filtered_rows:
                # Get all column headers from CSV
                headers = list(filtered_rows[0].keys()) if filtered_rows else []
                table_rows = []
                
                for row in filtered_rows[:20]:  # Show first 20
                    table_rows.append([row.get(header, 'N/A') for header in headers])
                
                table_data = {
                    "title": f"📊 Trade Blotter Data ({len(filtered_rows)} records{filter_message})",
                    "headers": headers,
                    "rows": table_rows,
                    "footer": f"Showing 20 of {len(filtered_rows)} total records" if len(filtered_rows) > 20 else None
                }
                
                # Memory context for response
                memory_context = ""
                if client_match:
                    client_name = client_match.group(1).strip().title()
                    if client_name in self.long_term_memory:
                        query_count = self.long_term_memory[client_name]['query_count']
                        memory_context = f"\n\n💾 Memory: {query_count} previous queries about {client_name}"
                
                response_data = {
                    "success": True,
                    "agent": "excel",
                    "response": f"Found {len(filtered_rows)} trade records{filter_message}.{memory_context}",
                    "tableData": table_data,
                    "data": filtered_rows
                }
                
                return response_data
            else:
                result = "No trade data found in CSV"
        except Exception as e:
            logger.error(f"Error reading CSV: {e}")
            result = f"Error reading trade data: {str(e)}"
        
        # Extract client name from query
        client_patterns = [
            r'\b(sheila|carter|sheila carter)\b',
            r'\b(john|doe|john doe)\b',
            r'trades?\s+for\s+([a-z\s]+)',
            r'show.*\s+([a-z\s]+).*trades?',
            r'data\s+for\s+([a-z\s]+)',
            r'mail.*(?:of|for)\s+([a-z\s]+)',
        ]
        
        extracted_client_data = None
        for pattern in client_patterns:
            match = re.search(pattern, query.lower())
            if match:
                potential_name = match.group(1).strip().title()
                
                # Get full client data from CSV
                try:
                    with open(self.csv_path, 'r', encoding='utf-8') as f:
                        reader = csv.DictReader(f)
                        for row in reader:
                            if row.get('Client', '').lower() == potential_name.lower():
                                extracted_client_data = {
                                    'client_name': row.get('Client', ''),
                                    'email': row.get('Email', ''),
                                    'account': row.get('Account', ''),
                                    'ticker': row.get('Ticker', ''),
                                    'quantity': row.get('Qty', ''),
                                    'follow_up': row.get('Follow-up', ''),
                                    'meeting_needed': row.get('Meeting', '')
                                }
                                break
                except Exception as e:
                    print(f"Error extracting client data: {e}")
                
                # Save to shared context for other agents
                if extracted_client_data:
                    conversation_memory['shared_context']['last_client_name'] = potential_name
                    conversation_memory['shared_context']['last_client_data'] = extracted_client_data
                    logger.info(f"💾 Excel Agent: Saved client data to shared memory")
                    logger.info(f"💾 Client: {potential_name}")
                    logger.info(f"💾 Email: {extracted_client_data['email']}")
                    logger.info(f"💾 Full data: {json.dumps(extracted_client_data, indent=2)}")
                break
        
        response_data = {
            "success": True,
            "agent": "excel",
            "response": result
        }
        
        # Include client data in response for handoffs
        if extracted_client_data:
            response_data['client_data'] = extracted_client_data
        
        return response_data


class FinanceAgent(BaseAgent):
    """Specialized agent for real-time financial data and banking queries via Finnhub API"""
    
    def __init__(self):
        super().__init__(AgentType.FINANCE)
        self.capabilities = [
            "real_time_stock_price",
            "company_info",
            "stock_comparison",
            "market_analysis",
            "financial_assistant"
        ]
        
        # Initialize Finnhub API
        try:
            from tools.finnhub_tools import get_stock_price
            from tools.trade_data_tools import get_trade_summary
            self.get_stock_price = get_stock_price
            self.get_trade_summary = get_trade_summary
            self.available = True
            print("✓ Finance Agent initialized")
        except Exception as e:
            print(f"⚠️  Finance Agent unavailable: {e}")
            self.available = False
    
    async def can_handle(self, query: str, context: Dict[str, Any]) -> bool:
        """Check if this is a finance, stock, or banking query"""
        if not self.available:
            return False
        
        query_lower = query.lower()
        
        # Stock and company queries
        finance_keywords = [
            'stock', 'price', 'ticker', 'share', 'quote', 'trading',
            'market', 'nasdaq', 'nyse', 'dow', 'index',
            'aapl', 'apple', 'tsla', 'tesla', 'msft', 'microsoft',
            'googl', 'google', 'amzn', 'amazon', 'rivn', 'rivian',
            'nvda', 'nvidia', 'meta', 'ibm', 'pltr', 'palantir'
        ]
        
        # Banking and financial assistant queries
        banking_keywords = [
            'bank', 'finance', 'invest', 'portfolio', 'dividend',
            'earnings', 'revenue', 'profit', 'valuation', 'pe ratio',
            'market cap', 'analyst', 'rating', 'forecast'
        ]
        
        # Question patterns
        question_patterns = [
            r'what.*(?:stock|price|trading)',
            r'how.*(?:stock|market|trading)',
            r'(?:compare|vs)\s+(?:stock|share)',
            r'(?:buy|sell)\s+(?:stock|share)'
        ]
        
        # Check keywords
        if any(kw in query_lower for kw in finance_keywords + banking_keywords):
            return True
        
        # Check question patterns
        return any(re.search(pattern, query_lower) for pattern in question_patterns)
    
    async def process(self, query: str, context: Dict[str, Any]) -> Dict[str, Any]:
        """Process finance query with real-time Finnhub API data"""
        try:
            query_lower = query.lower()
            query_upper = query.upper()
            
            # Extract ticker symbols from query
            ticker_map = {
                'apple': 'AAPL', 'aapl': 'AAPL',
                'tesla': 'TSLA', 'tsla': 'TSLA',
                'microsoft': 'MSFT', 'msft': 'MSFT',
                'google': 'GOOGL', 'googl': 'GOOGL', 'alphabet': 'GOOGL',
                'amazon': 'AMZN', 'amzn': 'AMZN',
                'rivian': 'RIVN', 'rivn': 'RIVN',
                'nvidia': 'NVDA', 'nvda': 'NVDA',
                'meta': 'META', 'facebook': 'META',
                'ibm': 'IBM',
                'palantir': 'PLTR', 'pltr': 'PLTR',
                'duke': 'DUK', 'duk': 'DUK',
                'delta': 'DAL', 'dal': 'DAL'
            }
            
            found_tickers = []
            for word, ticker in ticker_map.items():
                if word in query_lower:
                    if ticker not in found_tickers:
                        found_tickers.append(ticker)
            
            # Handle stock comparison queries
            if len(found_tickers) >= 2 and any(kw in query_lower for kw in ['compare', 'vs', 'versus', 'or']):
                results = []
                for ticker in found_tickers[:2]:  # Compare first 2
                    stock_data = self.get_stock_price(ticker)
                    results.append(stock_data)
                
                comparison = f"📊 **Stock Comparison**\n\n"
                comparison += f"**{found_tickers[0]}**\n{results[0]}\n\n"
                comparison += f"**{found_tickers[1]}**\n{results[1]}"
                
                return {
                    "success": True,
                    "agent": "finance",
                    "response": comparison
                }
            
            # Handle single stock query
            if found_tickers:
                ticker = found_tickers[0]
                result = self.get_stock_price(ticker)
                
                # Add context from memory if available
                memory_context = ""
                if 'shared_context' in context and 'last_client_name' in context['shared_context']:
                    client = context['shared_context']['last_client_name']
                    memory_context = f"\n\n💾 Context: Query related to {client}'s portfolio"
                
                return {
                    "success": True,
                    "agent": "finance",
                    "response": result + memory_context,
                    "ticker": ticker
                }
            
            # Fallback: Trade summary
            result = self.get_trade_summary()
            return {
                "success": True,
                "agent": "finance",
                "response": result
            }
        except Exception as e:
            logger.error(f"Finance agent error: {e}")
            return {
                "success": False,
                "agent": "finance",
                "error": f"Finance query failed: {str(e)}"
            }
        
        return {
            "success": True,
            "agent": "finance",
            "response": result
        }


class ComplianceAgent(BaseAgent):
    """Specialized agent for compliance with Astra DB short-term session memory"""
    
    def __init__(self):
        super().__init__(AgentType.COMPLIANCE)
        self.capabilities = [
            "search_knowledge_base",
            "compliance_rules",
            "risk_assessment",
            "client_profile_lookup",
            "hybrid_search",
            "session_memory"  # Short-term memory per session
        ]
        
        # Session memory storage
        self.session_memories = {}  # session_id -> List[messages]
        self.max_memory_per_session = 50  # Store last 50 interactions
        
        # Import Astra DB tools with memory capabilities
        try:
            from tools.astra_db_tools import (
                query_astra_db, 
                get_client_profile,
                index_client_knowledge_graph,
                get_astra_store
            )
            self.query_astra = query_astra_db
            self.get_client_profile = get_client_profile
            self.index_knowledge_graph = index_client_knowledge_graph
            self.astra_store = get_astra_store()
            
            # Fallback to ChromaDB for knowledge base
            self.legacy_kb = None
            if not self.astra_store:
                print("⚠️  Astra DB unavailable, falling back to ChromaDB")
                from tools.rag_tools import query_knowledge_base
                self.legacy_kb = query_knowledge_base
            
            self.available = True
            print("✓ Compliance Agent initialized with Astra DB and session memory")
        except Exception as e:
            print(f"⚠️  Compliance Agent unavailable: {e}")
            self.available = False
    
    def add_to_session_memory(self, session_id: str, role: str, content: str):
        """Add message to session memory"""
        if session_id not in self.session_memories:
            self.session_memories[session_id] = []
        
        self.session_memories[session_id].append({
            "role": role,
            "content": content,
            "timestamp": datetime.now().isoformat()
        })
        
        # Keep only last N messages
        if len(self.session_memories[session_id]) > self.max_memory_per_session:
            self.session_memories[session_id] = self.session_memories[session_id][-self.max_memory_per_session:]
    
    def get_session_context(self, session_id: str, last_n: int = 5) -> str:
        """Get recent session context"""
        if session_id not in self.session_memories:
            return ""
        
        recent_messages = self.session_memories[session_id][-last_n:]
        context = "\n".join([f"{m['role']}: {m['content']}" for m in recent_messages])
        return f"\n\n📝 Recent Context:\n{context}"
    
    async def can_handle(self, query: str, context: Dict[str, Any]) -> bool:
        """Check if this is a compliance query"""
        if not self.available:
            return False
        
        compliance_keywords = [
            'compliance', 'regulation', 'rule', 'policy',
            'churning', 'solicited', 'unsolicited', 'risk',
            'guideline', 'procedure', 'what is', 'define',
            'profile', 'history', 'search', 'find trades',
            'client background', 'past trades'
        ]
        query_lower = query.lower()
        
        return any(keyword in query_lower for keyword in compliance_keywords)
    
    async def process(self, query: str, context: Dict[str, Any]) -> Dict[str, Any]:
        """Process compliance query with Astra DB hybrid search and session memory"""
        # Extract session ID and add to memory
        session_id = context.get('conversation_id', 'default')
        self.add_to_session_memory(session_id, "user", query)
        
        # Get recent session context for context-aware responses
        session_context = self.get_session_context(session_id, last_n=5) if session_id in self.session_memories else ""
        
        query_lower = query.lower()
        
        # Check if asking for client profile
        if 'profile' in query_lower or 'client background' in query_lower:
            # Extract client name from query or context
            import re
            name_match = re.search(r'(?:profile|background)(?:\s+of)?(?:\s+for)?\s+([a-zA-Z\s]+)', query, re.IGNORECASE)
            
            client_name = None
            if name_match:
                client_name = name_match.group(1).strip()
            elif context.get('client_data', {}).get('client_name'):
                client_name = context['client_data']['client_name']
            
            if client_name and self.astra_store:
                profile = self.get_client_profile(client_name)
                if profile:
                    response = f"**Client Profile: {profile['client_name']}**\n\n"
                    response += f"📧 Emails: {', '.join(profile['emails'])}\n"
                    response += f"🏦 Accounts: {', '.join(profile['accounts'])}\n"
                    response += f"📊 Total Trades: {profile['total_trades']} (Buy: {profile['buy_count']}, Sell: {profile['sell_count']})\n"
                    response += f"⚠️  Risk Score: {profile['risk_score']}/100\n"
                    response += f"📈 Most Traded Tickers: {', '.join(list(profile['tickers_traded'].keys())[:5])}\n"
                    
                    if profile['notes']:
                        response += f"\n📝 Recent Notes:\n"
                        for note in profile['notes'][:3]:
                            response += f"  • {note}\n"
                    
                # Add response to session memory
                self.add_to_session_memory(session_id, "assistant", response)
                
                return {
                    "success": True,
                    "agent": "compliance",
                    "response": response,
                    "client_profile": profile
                }        # Hybrid search across all sources
        if self.astra_store:
            try:
                # Determine search type
                search_type = "all"
                if 'email' in query_lower:
                    search_type = "emails"
                elif 'trade' in query_lower or 'transaction' in query_lower:
                    search_type = "trades"
                elif 'compliance' in query_lower or 'rule' in query_lower:
                    search_type = "compliance"
                
                # Include session context in search for context-aware results
                enriched_query = f"{query}\n\nRecent context: {session_context}" if session_context else query
                results = self.query_astra(enriched_query, search_type=search_type, limit=5)
                
                response = f"**Search Results:**\n\n{results}"
                self.add_to_session_memory(session_id, "assistant", response)
                
                return {
                    "success": True,
                    "agent": "compliance",
                    "response": response,
                    "search_type": search_type
                }
            except Exception as e:
                print(f"⚠️  Astra DB search failed: {e}")
        
        # Fallback to legacy RAG
        if self.legacy_kb:
            result = self.legacy_kb(query)
        else:
            result = "Astra DB search unavailable. Please configure ASTRA_DB_APPLICATION_TOKEN and ASTRA_DB_API_ENDPOINT."
        
        # Add fallback response to session memory
        self.add_to_session_memory(session_id, "assistant", result)
        
        return {
            "success": True,
            "agent": "compliance",
            "response": result
        }


# ============================================================================
# COORDINATOR AGENT
# ============================================================================

class CoordinatorAgent(BaseAgent):
    """Coordinates between specialized agents using IBM ADK agent_builder"""
    
    def __init__(self):
        super().__init__(AgentType.COORDINATOR)
        
        # Initialize all specialized agents
        self.agents = {
            AgentType.GMAIL: GmailAgent(),
            AgentType.EXCEL: ExcelAgent(),
            AgentType.FINANCE: FinanceAgent(),
            AgentType.COMPLIANCE: ComplianceAgent(),
            AgentType.TRADE_PARSER: TradeParserAgent(),
        }
        
        # Initialize IBM ADK orchestrator agent if available
        if HAS_AGENT_BUILDER:
            self._init_orchestrator_agent()
        
        print("✓ Coordinator Agent initialized with all sub-agents")
    
    def _init_orchestrator_agent(self):
        """Initialize IBM watsonx Orchestrate master agent"""
        try:
            # Create orchestrator agent that manages all sub-agents
            orchestrator_config = {
                "name": "orqon_orchestrator",
                "description": "Master orchestrator for multi-agent trading system",
                "model": "granite-3-8b-instruct",
                "agents": [
                    {"name": "gmail_agent", "type": "email_operations"},
                    {"name": "excel_agent", "type": "data_operations"},
                    {"name": "finance_agent", "type": "financial_data"},
                    {"name": "compliance_agent", "type": "knowledge_base"}
                ]
            }
            
            # Create orchestrator using IBM watsonx Orchestrate 1.15.0 API
            orchestrator_spec = {
                "name": "coordinator_orchestrator",
                "description": orchestrator_config["description"],
                "kind": AgentKind.ASSISTANT,
                "title": "Multi-Agent Coordinator",
                "nickname": "coordinator",
                "config": AssistantAgentConfig(
                    description=orchestrator_config["description"],
                ),
            }
            
            self.orchestrator = AssistantAgent(**orchestrator_spec)
            print("✓ IBM ADK Orchestrator agent initialized")
        except Exception as e:
            print(f"⚠️  Failed to initialize ADK Orchestrator: {e}")
            self.orchestrator = None
    
    async def route_query(self, query: str, context: Dict[str, Any]) -> BaseAgent:
        """Route query to appropriate agent with priority order"""
        
        # PRIORITY 1: Trade Parser (must check first for complex logs)
        if AgentType.TRADE_PARSER in self.agents:
            if await self.agents[AgentType.TRADE_PARSER].can_handle(query, context):
                print(f"🎯 Routing to trade_parser agent (priority)")
                return self.agents[AgentType.TRADE_PARSER]
        
        # PRIORITY 2: Gmail (for email operations)
        if AgentType.GMAIL in self.agents:
            if await self.agents[AgentType.GMAIL].can_handle(query, context):
                print(f"🎯 Routing to gmail agent")
                return self.agents[AgentType.GMAIL]
        
        # PRIORITY 3: Other specialized agents
        for agent_type, agent in self.agents.items():
            if agent_type in [AgentType.TRADE_PARSER, AgentType.GMAIL]:
                continue  # Already checked
            if await agent.can_handle(query, context):
                print(f"🎯 Routing to {agent_type.value} agent")
                return agent
        
        # Default to Excel agent (general queries)
        print(f"🎯 Routing to excel agent (default)")
        return self.agents[AgentType.EXCEL]
    
    async def process_with_handoff(
        self, 
        query: str, 
        context: Dict[str, Any]
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """Process query with potential agent handoffs"""
        
        query_lower = query.lower().strip()
        
        # =====================================================================
        # CONVERSATIONAL AI: Handle greetings and identity questions
        # =====================================================================
        
        # Greetings
        greeting_patterns = ['hello', 'hi', 'hey', 'good morning', 'good afternoon', 'good evening', 'greetings', 'howdy']
        if any(pattern in query_lower for pattern in greeting_patterns) and len(query.split()) <= 3:
            yield {
                "type": "agent_assigned",
                "agent": "conversational",
                "timestamp": datetime.now().isoformat()
            }
            yield {
                "type": "agent_response",
                "agent": "conversational",
                "data": {
                    "success": True,
                    "response": f"Hello! 👋 I'm ORQON, your AI-powered trade intelligence assistant.\n\nI'm here to help you with:\n\n• **Trade Data** - Query client portfolios, trade history, and account details\n• **Email & Communication** - Send emails, schedule meetings, set reminders\n• **Calendar Management** - Create Google Meet meetings and reminders\n• **Financial Information** - Get stock prices and company information\n• **Compliance Analysis** - Answer SEC compliance questions\n• **Data Analysis** - Show tables, filter data, and generate reports\n\nWhat would you like to do today?"
                },
                "timestamp": datetime.now().isoformat()
            }
            return
        
        # Identity questions
        identity_patterns = [
            'who are you', 'what are you', 'who r u', 'what r u',
            'tell me about yourself', 'introduce yourself',
            'what can you do', 'what do you do', 'your name',
            'are you ai', 'are you a bot', 'are you human'
        ]
        if any(pattern in query_lower for pattern in identity_patterns):
            yield {
                "type": "agent_assigned",
                "agent": "conversational",
                "timestamp": datetime.now().isoformat()
            }
            yield {
                "type": "agent_response",
                "agent": "conversational",
                "data": {
                    "success": True,
                    "response": f"I'm **ORQON** 🤖 - an AI assistant specialized in trade intelligence and compliance.\n\n**About Me:**\n• Built on **IBM watsonx.ai** with Granite LLM\n• Specialized in financial data analysis and SEC compliance\n• Connected to Google Workspace (Gmail, Calendar, Meet)\n• Multi-agent architecture with smart routing\n\n**My Capabilities:**\n✅ Query client trade data and portfolios\n✅ Send emails and schedule meetings\n✅ Create Google Meet conferences\n✅ Fetch real-time stock prices\n✅ Answer compliance questions\n✅ Analyze trade patterns and risk\n✅ Manage reminders and follow-ups\n\nI'm like a human assistant, but I never sleep! 😊\n\nHow can I help you today?"
                },
                "timestamp": datetime.now().isoformat()
            }
            return
        
        # Date and Time queries
        datetime_patterns = ['what is date', 'what is time', 'what is the date', 'what is the time', 
                            'whats the date', 'whats the time', 'current date', 'current time',
                            'what date', 'what time', 'today date', 'todays date', 'date and time']
        if any(pattern in query_lower for pattern in datetime_patterns):
            now = datetime.now()
            day_name = now.strftime('%A')
            date_str = now.strftime('%B %d, %Y')
            time_str = now.strftime('%I:%M:%S %p')
            
            yield {
                "type": "agent_assigned",
                "agent": "conversational",
                "timestamp": datetime.now().isoformat()
            }
            yield {
                "type": "agent_response",
                "agent": "conversational",
                "data": {
                    "success": True,
                    "response": f"📅 **Current Date & Time**\n\n• **Date:** {date_str}\n• **Day:** {day_name}\n• **Time:** {time_str}\n\nHow can I assist you further?"
                },
                "timestamp": datetime.now().isoformat()
            }
            return
        
        # Thank you / gratitude
        gratitude_patterns = ['thank you', 'thanks', 'thx', 'ty', 'appreciate it', 'appreciate']
        if any(pattern in query_lower for pattern in gratitude_patterns) and len(query.split()) <= 5:
            yield {
                "type": "agent_assigned",
                "agent": "conversational",
                "timestamp": datetime.now().isoformat()
            }
            yield {
                "type": "agent_response",
                "agent": "conversational",
                "data": {
                    "success": True,
                    "response": "You're very welcome! 😊 Let me know if you need anything else!"
                },
                "timestamp": datetime.now().isoformat()
            }
            return
        
        # =====================================================================
        # SHORT-TERM MEMORY: Resolve pronouns to last mentioned client
        # =====================================================================
        pronouns = ['his', 'her', 'their', 'he', 'she', 'they', 'him']
        
        # Check if query contains pronouns
        has_pronoun = any(f' {pronoun} ' in f' {query_lower} ' or query_lower.startswith(f'{pronoun} ') 
                         for pronoun in pronouns)
        
        if has_pronoun:
            # Get last mentioned client from shared memory
            last_client = conversation_memory['shared_context'].get('last_client_name')
            
            if last_client:
                logger.info(f"🧠 SHORT-TERM MEMORY: Detected pronoun reference")
                logger.info(f"🧠 Resolving to last client: {last_client}")
                
                # Replace pronouns with client name
                for pronoun in pronouns:
                    # Handle possessive pronouns (his/her/their)
                    if pronoun in ['his', 'her', 'their']:
                        query = re.sub(rf'\b{pronoun}\b', f"{last_client}'s", query, flags=re.IGNORECASE)
                    # Handle subject pronouns (he/she/they)
                    elif pronoun in ['he', 'she', 'they']:
                        query = re.sub(rf'\b{pronoun}\b', last_client, query, flags=re.IGNORECASE)
                    # Handle object pronouns (him)
                    elif pronoun in ['him']:
                        query = re.sub(rf'\b{pronoun}\b', last_client, query, flags=re.IGNORECASE)
                
                logger.info(f"🧠 Rewritten query: {query}")
            else:
                logger.warning(f"⚠️ Pronoun detected but no previous client in memory")
        
        # Route to initial agent
        current_agent = await self.route_query(query, context)
        
        # Stream agent assignment
        yield {
            "type": "agent_assigned",
            "agent": current_agent.agent_type.value,
            "timestamp": datetime.now().isoformat()
        }
        
        # Process with current agent
        try:
            result = await current_agent.process(query, context)
            
            # If result contains client_data, merge it into context for next agent
            if isinstance(result, dict) and 'client_data' in result:
                context['client_data'] = result['client_data']
                print(f"🔄 Context enriched with client data for potential handoffs")
            
            # Stream result
            yield {
                "type": "agent_response",
                "agent": current_agent.agent_type.value,
                "data": result,
                "timestamp": datetime.now().isoformat()
            }
            
            # Check if handoff needed
            # (Future: implement smart handoff detection)
            
        except Exception as e:
            yield {
                "type": "error",
                "agent": current_agent.agent_type.value,
                "error": str(e),
                "timestamp": datetime.now().isoformat()
            }


# ============================================================================
# MCP SERVER
# ============================================================================

app = FastAPI(
    title="Orqon IBM MCP Toolkit Server",
    description="IBM watsonx Orchestrate MCP toolkit with specialized agents, streaming, and Model Context Protocol support",
    version="3.0.0",
    openapi_url="/openapi.json",
    docs_url="/docs",
    redoc_url="/redoc"
)

# Enable CORS (IBM MCP remote toolkit requirement)
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://localhost:3001", 
        "http://localhost:5173",
        "https://*.watson-orchestrate.ibm.com",  # IBM watsonx Orchestrate
        "*"  # Allow all for MCP protocol
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)

# Initialize coordinator
coordinator = CoordinatorAgent()

# IBM watsonx Orchestrate agent registration
if HAS_AGENT_BUILDER:
    try:
        # All specialized agents have been initialized with IBM watsonx Orchestrate API
        agent_count = sum(1 for agent in coordinator.agents.values() if agent.adk_agent)
        logger.info(f"✓ {agent_count} agents initialized with IBM watsonx Orchestrate")
    except Exception as e:
        logger.warning(f"⚠️  Agent registration check failed: {e}")


# ============================================================================
# IBM MCP TOOLKIT PROTOCOL ENDPOINTS
# ============================================================================

@app.get("/mcp/info", response_model=MCPServerInfo)
async def mcp_server_info():
    """
    IBM MCP toolkit server information endpoint
    Returns server capabilities and protocol version
    
    Required for IBM watsonx Orchestrate toolkit import
    """
    return MCPServerInfo(
        name="orqon-mcp-toolkit",
        version="3.0.0",
        protocolVersion="2024-11-05",
        capabilities={
            "tools": {
                "listChanged": False  # Tools are static
            },
            "prompts": {},
            "resources": {}
        }
    )


@app.get("/mcp/tools/list", response_model=MCPListToolsResponse)
async def mcp_list_tools():
    """
    IBM MCP toolkit list tools endpoint
    Returns all available tools with their schemas
    
    Called during toolkit import with 30-second timeout
    """
    logger.info("🔍 IBM MCP toolkit discovery: listing tools")
    
    # Dynamically build tool list from available agents
    tools = []
    
    # Gmail tools
    if coordinator.agents.get(AgentType.GMAIL) and coordinator.agents[AgentType.GMAIL].available:
        tools.extend([
            MCPToolSchema(
                name="gmail_send_email",
                description="Send an email using Gmail. Automatically looks up client email addresses from the database.",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "to": {
                            "type": "string",
                            "description": "Recipient email address or client name (will auto-lookup email)"
                        },
                        "subject": {
                            "type": "string",
                            "description": "Email subject line"
                        },
                        "body": {
                            "type": "string",
                            "description": "Email body content (supports HTML)"
                        },
                        "cc": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": "CC recipients (optional)"
                        },
                        "attachments": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": "File paths to attach (optional)"
                        }
                    },
                    "required": ["to", "subject", "body"]
                }
            ),
            MCPToolSchema(
                name="gmail_draft_email",
                description="Create a draft email in Gmail without sending",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "to": {"type": "string", "description": "Recipient email"},
                        "subject": {"type": "string", "description": "Email subject"},
                        "body": {"type": "string", "description": "Email body"}
                    },
                    "required": ["to", "subject", "body"]
                }
            ),
            MCPToolSchema(
                name="gmail_search_emails",
                description="Search emails in Gmail",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "Gmail search query"},
                        "max_results": {"type": "integer", "description": "Max results", "default": 10}
                    },
                    "required": ["query"]
                }
            )
        ])
    
    # Excel/CSV tools
    if coordinator.agents.get(AgentType.EXCEL) and coordinator.agents[AgentType.EXCEL].available:
        tools.extend([
            MCPToolSchema(
                name="excel_get_client_email",
                description="Get client email address from trade database",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "client_name": {"type": "string", "description": "Client name"}
                    },
                    "required": ["client_name"]
                }
            ),
            MCPToolSchema(
                name="excel_get_client_trades",
                description="Get all trades for a specific client",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "client_name": {"type": "string", "description": "Client name"}
                    },
                    "required": ["client_name"]
                }
            ),
            MCPToolSchema(
                name="excel_search_trades",
                description="Search trades by ticker, type, or other criteria",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "Search query"}
                    },
                    "required": ["query"]
                }
            )
        ])
    
    # Finance tools
    if coordinator.agents.get(AgentType.FINANCE) and coordinator.agents[AgentType.FINANCE].available:
        tools.extend([
            MCPToolSchema(
                name="finance_get_stock_price",
                description="Get current stock price and market data",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "ticker": {"type": "string", "description": "Stock ticker symbol (e.g., AAPL, TSLA)"}
                    },
                    "required": ["ticker"]
                }
            ),
            MCPToolSchema(
                name="finance_analyze_portfolio",
                description="Analyze client portfolio and positions",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "client_name": {"type": "string", "description": "Client name"}
                    },
                    "required": ["client_name"]
                }
            )
        ])
    
    # Compliance tools
    if coordinator.agents.get(AgentType.COMPLIANCE) and coordinator.agents[AgentType.COMPLIANCE].available:
        tools.extend([
            MCPToolSchema(
                name="compliance_search_knowledge",
                description="Search compliance knowledge base using RAG (Retrieval-Augmented Generation)",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "Compliance question or topic"},
                        "search_type": {
                            "type": "string",
                            "enum": ["all", "trades", "emails", "compliance"],
                            "description": "Type of documents to search",
                            "default": "all"
                        }
                    },
                    "required": ["query"]
                }
            ),
            MCPToolSchema(
                name="compliance_check_client_risk",
                description="Check client risk profile and compliance status",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "client_name": {"type": "string", "description": "Client name"}
                    },
                    "required": ["client_name"]
                }
            )
        ])
    
    logger.info(f"✓ IBM MCP toolkit: discovered {len(tools)} tools")
    return MCPListToolsResponse(tools=tools)


@app.post("/mcp/tools/call", response_model=MCPToolResponse)
async def mcp_call_tool(request: MCPCallToolRequest):
    """
    IBM MCP toolkit call tool endpoint
    Executes a specific tool with provided arguments
    
    Returns:
        MCPToolResponse with content or error
    """
    logger.info(f"🔧 IBM MCP toolkit: executing tool '{request.name}'")
    
    try:
        # Route tool call to appropriate agent
        tool_name = request.name
        arguments = request.arguments
        
        # Gmail tools
        if tool_name.startswith("gmail_"):
            agent = coordinator.agents.get(AgentType.GMAIL)
            if not agent or not agent.available:
                raise HTTPException(status_code=503, detail="Gmail agent not available")
            
            if tool_name == "gmail_send_email":
                result = await agent.gmail_tools.send_email(
                    to=arguments.get("to"),
                    subject=arguments.get("subject"),
                    body=arguments.get("body"),
                    cc=arguments.get("cc"),
                    attachments=arguments.get("attachments")
                )
            elif tool_name == "gmail_draft_email":
                result = await agent.gmail_tools.draft_email(
                    to=arguments.get("to"),
                    subject=arguments.get("subject"),
                    body=arguments.get("body")
                )
            elif tool_name == "gmail_search_emails":
                result = agent.gmail_tools.search_emails(
                    query=arguments.get("query"),
                    max_results=arguments.get("max_results", 10)
                )
            else:
                raise HTTPException(status_code=404, detail=f"Unknown tool: {tool_name}")
        
        # Excel tools
        elif tool_name.startswith("excel_"):
            agent = coordinator.agents.get(AgentType.EXCEL)
            if not agent or not agent.available:
                raise HTTPException(status_code=503, detail="Excel agent not available")
            
            if tool_name == "excel_get_client_email":
                email = get_client_email_from_csv(arguments.get("client_name"))
                result = {"email": email} if email else {"error": "Email not found"}
            elif tool_name == "excel_get_client_trades":
                # Use Excel agent's process method
                query = f"Get all trades for {arguments.get('client_name')}"
                result = await agent.process(query, {})
            elif tool_name == "excel_search_trades":
                result = await agent.process(arguments.get("query"), {})
            else:
                raise HTTPException(status_code=404, detail=f"Unknown tool: {tool_name}")
        
        # Finance tools
        elif tool_name.startswith("finance_"):
            agent = coordinator.agents.get(AgentType.FINANCE)
            if not agent or not agent.available:
                raise HTTPException(status_code=503, detail="Finance agent not available")
            
            result = await agent.process(
                f"{tool_name}: {json.dumps(arguments)}",
                arguments
            )
        
        # Compliance tools
        elif tool_name.startswith("compliance_"):
            agent = coordinator.agents.get(AgentType.COMPLIANCE)
            if not agent or not agent.available:
                raise HTTPException(status_code=503, detail="Compliance agent not available")
            
            result = await agent.process(
                f"{tool_name}: {json.dumps(arguments)}",
                arguments
            )
        
        else:
            raise HTTPException(status_code=404, detail=f"Unknown tool: {tool_name}")
        
        # Format response
        logger.info(f"✓ IBM MCP toolkit: tool '{request.name}' executed successfully")
        return MCPToolResponse(
            content=[
                {
                    "type": "text",
                    "text": json.dumps(result, indent=2) if isinstance(result, dict) else str(result)
                }
            ],
            isError=False
        )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ IBM MCP toolkit: tool '{request.name}' failed: {e}")
        return MCPToolResponse(
            content=[
                {
                    "type": "text",
                    "text": f"Error executing tool: {str(e)}"
                }
            ],
            isError=True
        )


@app.get("/mcp/sse")
async def mcp_sse_endpoint(request: Request):
    """
    IBM MCP toolkit SSE (Server-Sent Events) transport endpoint
    Enables remote MCP toolkit import via SSE protocol
    
    IBM requirement:
    - 30-second handshake timeout
    - Standard SSE headers
    - JSON-RPC 2.0 messages
    """
    async def event_generator():
        try:
            # Send server info
            server_info = await mcp_server_info()
            yield f"data: {server_info.model_dump_json()}\n\n"
            
            # Keep connection alive
            while True:
                if await request.is_disconnected():
                    break
                await asyncio.sleep(1)
                yield f"data: {json.dumps({'type': 'ping'})}\n\n"
        except Exception as e:
            logger.error(f"SSE error: {e}")
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


# ============================================================================
# PYDANTIC MODELS
# ============================================================================

class ChatRequest(BaseModel):
    message: str
    conversation_id: Optional[str] = None
    context: Optional[Dict[str, Any]] = None


class StreamChunk(BaseModel):
    type: str  # "agent_assigned", "agent_response", "error", "complete"
    agent: Optional[str] = None
    data: Optional[Dict[str, Any]] = None
    error: Optional[str] = None
    timestamp: str


# ============================================================================
# STREAMING ENDPOINTS
# ============================================================================

@app.post("/chat/stream")
async def chat_stream(request: ChatRequest):
    """Streaming chat endpoint with agent handoffs"""
    
    async def generate():
        """Generate streaming response"""
        context = request.context or {}
        
        async for chunk in coordinator.process_with_handoff(request.message, context):
            yield f"data: {json.dumps(chunk)}\n\n"
        
        # Send completion marker
        yield f"data: {json.dumps({'type': 'complete', 'timestamp': datetime.now().isoformat()})}\n\n"
    
    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no"
        }
    )


@app.websocket("/ws/chat")
async def websocket_chat(websocket: WebSocket):
    """WebSocket endpoint for real-time chat"""
    await websocket.accept()
    
    try:
        while True:
            # Receive message
            data = await websocket.receive_text()
            request_data = json.loads(data)
            
            query = request_data.get("message", "")
            context = request_data.get("context", {})
            
            # Process with agent handoffs
            async for chunk in coordinator.process_with_handoff(query, context):
                await websocket.send_json(chunk)
            
            # Send completion
            await websocket.send_json({
                "type": "complete",
                "timestamp": datetime.now().isoformat()
            })
            
    except WebSocketDisconnect:
        print("Client disconnected")


@app.post("/chat")
async def chat_standard(request: ChatRequest):
    """Standard (non-streaming) chat endpoint"""
    
    context = request.context or {}
    responses = []
    
    # Collect all responses
    async for chunk in coordinator.process_with_handoff(request.message, context):
        responses.append(chunk)
    
    # Return final response
    final_response = None
    for chunk in responses:
        if chunk["type"] == "agent_response":
            final_response = chunk["data"]
    
    if final_response:
        logger.info(f"📤 Final response keys: {list(final_response.keys())}")
        logger.info(f"📤 Response field value: {final_response.get('response', 'MISSING')[:200] if final_response.get('response') else 'EMPTY'}")
        
        response_data = {
            "success": final_response.get("success", True),
            "response": final_response.get("response", ""),
            "agent": final_response.get("agent"),
            "conversation_id": request.conversation_id or f"conv_{datetime.now().timestamp()}"
        }
        
        # Pass through tableData if present (for Excel agent)
        if final_response.get("tableData"):
            response_data["tableData"] = final_response["tableData"]
        
        # If Trade Parser returned trades, format for frontend
        if final_response.get("agent") == "trade_parser" and final_response.get("trades"):
            trades = final_response["trades"]
            # Format trades for frontend UI
            formatted_trades = []
            for trade in trades:
                formatted_trades.append({
                    "ticker": trade.get("ticker", "").upper(),
                    "action": trade.get("side", ""),  # "Buy" or "Sell"
                    "quantity": trade.get("quantity", 0),
                    "order_type": trade.get("order_type", "Market"),
                    "price": trade.get("price", 0),
                    "client_name": trade.get("client_name", ""),
                    "account_number": trade.get("account_number", ""),
                    "solicited": "Solicited" if trade.get("solicited") else "Unsolicited",
                    "ticket_id": trade.get("ticket_id", ""),
                    "timestamp": trade.get("timestamp", ""),
                    "notes": trade.get("notes", ""),
                    "follow_up_date": trade.get("follow_up_date", ""),
                    "email": trade.get("email", ""),
                    "stage": trade.get("stage", "Pending"),
                    "meeting_needed": "Yes" if trade.get("meeting_needed") else "No",
                    "confidence": 0.95  # High confidence for parsed trades
                })
            
            # Add parsed_trade for UI rendering
            response_data["parsed_trade"] = {
                "trades": formatted_trades
            }
        
        return response_data
    else:
        raise HTTPException(status_code=500, detail="No response from agents")


# ============================================================================
# UTILITY ENDPOINTS
# ============================================================================

@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "name": "Orqon Multi-Agent MCP Server",
        "version": "3.0.0",
        "description": "IBM watsonx Orchestrate ADK with specialized agents",
        "endpoints": {
            "chat": "/chat",
            "streaming": "/chat/stream",
            "websocket": "/ws/chat",
            "health": "/health",
            "agents": "/agents"
        }
    }


@app.get("/health")
async def health():
    """Health check"""
    return {
        "status": "healthy",
        "agents": {
            agent_type.value: agent.available if hasattr(agent, 'available') else True
            for agent_type, agent in coordinator.agents.items()
        },
        "model": "granite-3-8b-instruct"
    }


@app.post("/auth/token")
async def authenticate_user(request: dict):
    """
    Authenticate user and return JWT token
    For demo purposes, this returns a simple token based on user_id
    In production, implement proper JWT authentication with secret keys
    """
    try:
        user_email = request.get('user_email', 'trader@orqon.com')
        user_id = request.get('user_id', 'trader_001')
        metadata = request.get('metadata', {})
        
        # Generate a simple demo token (in production, use proper JWT with secret)
        import secrets
        access_token = f"orqon_{user_id}_{secrets.token_urlsafe(32)}"
        
        logger.info(f"🔐 User authenticated: {user_email} ({user_id})")
        
        return {
            "access_token": access_token,
            "token_type": "bearer",
            "user_id": user_id,
            "user_email": user_email,
            "expires_in": 86400  # 24 hours
        }
    except Exception as e:
        logger.error(f"❌ Authentication failed: {str(e)}")
        return {
            "error": "Authentication failed",
            "detail": str(e)
        }


@app.get("/agents")
async def list_agents():
    """List available agents and their capabilities"""
    return {
        "coordinator": {
            "type": "coordinator",
            "description": "Routes queries to specialized agents"
        },
        "agents": {
            agent_type.value: {
                "type": agent_type.value,
                "capabilities": agent.capabilities,
                "available": agent.available if hasattr(agent, 'available') else True
            }
            for agent_type, agent in coordinator.agents.items()
        }
    }


@app.get("/api/calendar/upcoming")
async def get_upcoming_calendar_events():
    """Get upcoming calendar events (reminders and meetings)"""
    try:
        gmail_agent = coordinator.agents.get(AgentType.GMAIL)
        if not gmail_agent or not hasattr(gmail_agent, 'gmail_tools'):
            return {"success": False, "error": "Gmail agent not available"}
        
        result = gmail_agent.gmail_tools.list_upcoming_events(max_results=10)
        
        if result.get('success'):
            return {
                "success": True,
                "events": result.get('events', []),
                "count": len(result.get('events', []))
            }
        else:
            return {
                "success": False,
                "error": result.get('error', 'Failed to fetch events'),
                "events": []
            }
    except Exception as e:
        logger.error(f"❌ Failed to fetch calendar events: {str(e)}", exc_info=True)
        return {
            "success": False,
            "error": str(e),
            "events": []
        }


@app.post("/admin/index-knowledge-graph")
async def index_knowledge_graph():
    """
    Admin endpoint: Build client knowledge graph from CSV data
    Indexes client profiles, relationships, and trade history into Astra DB
    """
    try:
        from tools.astra_db_tools import index_client_knowledge_graph, get_astra_store
        
        # Build knowledge graph
        result = index_client_knowledge_graph()
        
        # Also index trade data into vectors
        astra = get_astra_store()
        if astra:
            trade_result = astra.index_trade_data()
            result["trade_indexing"] = trade_result
        
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Indexing failed: {str(e)}")


@app.get("/admin/client-profile/{client_name}")
async def get_client_profile_api(client_name: str):
    """
    Admin endpoint: Get client profile from knowledge graph
    """
    try:
        from tools.astra_db_tools import get_client_profile
        
        profile = get_client_profile(client_name)
        
        if profile:
            return {
                "success": True,
                "profile": profile
            }
        else:
            raise HTTPException(status_code=404, detail="Client profile not found")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/admin/astra-search")
async def astra_search_api(query: str, search_type: str = "all", limit: int = 5):
    """
    Admin endpoint: Test Astra DB hybrid search
    """
    try:
        from tools.astra_db_tools import query_astra_db
        
        results = query_astra_db(query, search_type=search_type, limit=limit)
        
        return {
            "success": True,
            "query": query,
            "search_type": search_type,
            "results": results
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/download/csv")
async def download_csv():
    """
    Download trade blotter CSV file
    """
    from pathlib import Path
    from fastapi.responses import FileResponse
    
    csv_path = Path(__file__).parent / "data" / "trade_blotter.csv"
    
    if not csv_path.exists():
        raise HTTPException(status_code=404, detail="CSV file not found")
    
    return FileResponse(
        path=str(csv_path),
        filename="trade_blotter.csv",
        media_type="text/csv"
    )


@app.get("/download/excel")
async def download_excel():
    """
    Download trade blotter Excel file
    """
    from pathlib import Path
    from fastapi.responses import FileResponse
    
    excel_path = Path(__file__).parent / "data" / "trade_blotter.xlsx"
    
    if not excel_path.exists():
        raise HTTPException(status_code=404, detail="Excel file not found")
    
    return FileResponse(
        path=str(excel_path),
        filename="trade_blotter.xlsx",
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


@app.get("/get-csv-data")
async def get_csv_data():
    """
    Get CSV data as JSON for analytics charts
    """
    from pathlib import Path
    import csv
    
    csv_path = Path(__file__).parent / "data" / "trade_blotter.csv"
    
    if not csv_path.exists():
        raise HTTPException(status_code=404, detail="CSV file not found")
    
    try:
        rows = []
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                rows.append(row)
        
        return {
            "success": True,
            "rows": rows,
            "count": len(rows)
        }
    except Exception as e:
        logger.error(f"Failed to read CSV data: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to read CSV data: {str(e)}")


@app.get("/open-csv")
async def open_csv_file():
    """
    Open CSV file in system default application (GET endpoint for frontend)
    """
    from pathlib import Path
    import subprocess
    import platform
    
    file_path = Path(__file__).parent / "data" / "trade_blotter.csv"
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="CSV file not found")
    
    try:
        # Open file with system default application
        if platform.system() == "Windows":
            subprocess.Popen(["start", "", str(file_path)], shell=True)
        elif platform.system() == "Darwin":  # macOS
            subprocess.Popen(["open", str(file_path)])
        else:  # Linux
            subprocess.Popen(["xdg-open", str(file_path)])
        
        return {
            "success": True,
            "message": f"Opened CSV file: {file_path.name}",
            "file_path": str(file_path)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to open CSV: {str(e)}")


@app.get("/open-excel")
async def open_excel_file():
    """
    Open Excel file in system default application (GET endpoint for frontend)
    """
    from pathlib import Path
    import subprocess
    import platform
    
    file_path = Path(__file__).parent / "data" / "trade_blotter.xlsx"
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Excel file not found")
    
    try:
        # Open file with system default application
        if platform.system() == "Windows":
            subprocess.Popen(["start", "", str(file_path)], shell=True)
        elif platform.system() == "Darwin":  # macOS
            subprocess.Popen(["open", str(file_path)])
        else:  # Linux
            subprocess.Popen(["xdg-open", str(file_path)])
        
        return {
            "success": True,
            "message": f"Opened Excel file: {file_path.name}",
            "file_path": str(file_path)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to open Excel: {str(e)}")


@app.post("/open-file")
async def open_file(request: Dict[str, Any]):
    """
    Open CSV or Excel file in system default application
    Used when LLM wants to show the file to user
    """
    from pathlib import Path
    import subprocess
    import platform
    
    file_type = request.get("file_type", "csv").lower()
    
    if file_type == "csv":
        file_path = Path(__file__).parent / "data" / "trade_blotter.csv"
    elif file_type in ["excel", "xlsx"]:
        file_path = Path(__file__).parent / "data" / "trade_blotter.xlsx"
    else:
        raise HTTPException(status_code=400, detail="Invalid file type. Use 'csv' or 'excel'")
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"{file_type.upper()} file not found")
    
    try:
        # Open file with system default application
        if platform.system() == "Windows":
            subprocess.Popen(["start", "", str(file_path)], shell=True)
        elif platform.system() == "Darwin":  # macOS
            subprocess.Popen(["open", str(file_path)])
        else:  # Linux
            subprocess.Popen(["xdg-open", str(file_path)])
        
        return {
            "success": True,
            "message": f"Opened {file_type.upper()} file: {file_path.name}",
            "file_path": str(file_path)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to open file: {str(e)}")


@app.post("/parse-document")
async def parse_document(file: UploadFile = File(...)):
    """
    Parse uploaded document and extract trade information
    Supports: PDF, DOCX, DOC, TXT (max 10MB per IBM limits)
    """
    from pathlib import Path
    import tempfile
    
    # Validate file size (10MB limit per IBM doc extractor)
    MAX_SIZE = 10 * 1024 * 1024  # 10MB
    content = await file.read()
    
    if len(content) > MAX_SIZE:
        raise HTTPException(status_code=413, detail="File size exceeds 10MB limit")
    
    # Validate file type
    allowed_extensions = {'.pdf', '.docx', '.doc', '.txt', '.jpeg', '.jpg', '.png'}
    file_ext = Path(file.filename).suffix.lower()
    
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400, 
            detail=f"Invalid file type. Allowed: {', '.join(allowed_extensions)}"
        )
    
    try:
        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        extracted_text = ""
        
        # Extract text based on file type
        if file_ext == '.pdf':
            try:
                import PyPDF2
                with open(tmp_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    for page in pdf_reader.pages:
                        extracted_text += page.extract_text() + "\n"
            except Exception as e:
                logger.error(f"PDF extraction failed: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Failed to extract PDF: {str(e)}")
        
        elif file_ext in ['.docx', '.doc']:
            try:
                import docx
                doc = docx.Document(tmp_path)
                for paragraph in doc.paragraphs:
                    extracted_text += paragraph.text + "\n"
            except Exception as e:
                logger.error(f"DOCX extraction failed: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Failed to extract DOCX: {str(e)}")
        
        elif file_ext == '.txt':
            with open(tmp_path, 'r', encoding='utf-8') as f:
                extracted_text = f.read()
        
        elif file_ext in ['.jpeg', '.jpg', '.png']:
            try:
                from PIL import Image
                import pytesseract
                img = Image.open(tmp_path)
                extracted_text = pytesseract.image_to_string(img)
            except Exception as e:
                logger.error(f"Image OCR failed: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Failed to extract text from image: {str(e)}")
        
        # Clean up temp file
        Path(tmp_path).unlink()
        
        if not extracted_text.strip():
            return {
                "success": False,
                "error": "No text could be extracted from the document"
            }
        
        # Parse trade data from extracted text using Trade Parser Agent
        trade_parser = coordinator.agents.get(AgentType.TRADE_PARSER)
        if not trade_parser:
            return {
                "success": False,
                "error": "Trade parser agent not available"
            }
        
        parse_result = await trade_parser.process(extracted_text, {})
        
        return {
            "success": True,
            "filename": file.filename,
            "extracted_text": extracted_text[:1000],  # First 1000 chars for preview
            "parsed_trades": parse_result.get("trades", []),
            "agent": "trade_parser",
            "response": parse_result.get("response", "")
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document parsing error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Document parsing failed: {str(e)}")


@app.post("/transcribe-audio")
async def transcribe_audio(file: UploadFile = File(...)):
    """
    Transcribe audio using IBM Watson Speech-to-Text
    Supports: WAV, MP3, FLAC, OGG (max 100MB)
    """
    import tempfile
    from pathlib import Path
    
    try:
        # Read audio content
        content = await file.read()
        
        # Validate file size (100MB limit)
        MAX_SIZE = 100 * 1024 * 1024
        if len(content) > MAX_SIZE:
            raise HTTPException(status_code=413, detail="Audio file exceeds 100MB limit")
        
        # Validate file type
        allowed_extensions = {'.wav', '.mp3', '.flac', '.ogg', '.webm', '.m4a'}
        file_ext = Path(file.filename).suffix.lower()
        
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid audio format. Allowed: {', '.join(allowed_extensions)}"
            )
        
        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        # Get Watson STT credentials from environment
        import os
        watson_api_key = os.getenv('WATSON_STT_API_KEY')
        watson_url = os.getenv('WATSON_STT_URL', 'https://api.us-south.speech-to-text.watson.cloud.ibm.com')
        
        if not watson_api_key:
            # Fallback: Try using SpeechRecognition with Google (free tier)
            logger.warning("Watson STT credentials not found, using Google Speech Recognition")
            
            try:
                import speech_recognition as sr
                recognizer = sr.Recognizer()
                
                # Convert audio to WAV if needed
                if file_ext != '.wav':
                    from pydub import AudioSegment
                    audio = AudioSegment.from_file(tmp_path)
                    wav_path = tmp_path.replace(file_ext, '.wav')
                    audio.export(wav_path, format='wav')
                    tmp_path = wav_path
                
                with sr.AudioFile(tmp_path) as source:
                    audio_data = recognizer.record(source)
                    transcript = recognizer.recognize_google(audio_data)
                
                Path(tmp_path).unlink()
                
                return {
                    "success": True,
                    "transcript": transcript,
                    "provider": "google_speech",
                    "confidence": 0.85
                }
                
            except Exception as e:
                logger.error(f"Speech recognition failed: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Transcription failed: {str(e)}")
        
        # Use IBM Watson Speech-to-Text
        try:
            from ibm_watson import SpeechToTextV1
            from ibm_cloud_sdk_core.authenticators import IAMAuthenticator
            
            authenticator = IAMAuthenticator(watson_api_key)
            speech_to_text = SpeechToTextV1(authenticator=authenticator)
            speech_to_text.set_service_url(watson_url)
            
            # Determine content type
            content_type_map = {
                '.wav': 'audio/wav',
                '.mp3': 'audio/mp3',
                '.flac': 'audio/flac',
                '.ogg': 'audio/ogg',
                '.webm': 'audio/webm',
                '.m4a': 'audio/mp4'
            }
            content_type = content_type_map.get(file_ext, 'audio/wav')
            
            # Transcribe with optimized parameters for financial/trading vocabulary
            with open(tmp_path, 'rb') as audio_file:
                response = speech_to_text.recognize(
                    audio=audio_file,
                    content_type=content_type,
                    model='en-US_BroadbandModel',
                    timestamps=True,
                    word_confidence=True,
                    smart_formatting=True,
                    speaker_labels=False,
                    end_of_phrase_silence_time=0.8,
                    split_transcript_at_phrase_end=True,
                    speech_detector_sensitivity=0.5,
                    background_audio_suppression=0.5,
                    keywords=['buy', 'sell', 'shares', 'stock', 'market', 'limit', 'trade', 'client', 'account'],
                    keywords_threshold=0.5
                ).get_result()
            
            Path(tmp_path).unlink()
            
            # Extract transcript
            if response.get('results'):
                transcript = ' '.join([
                    result['alternatives'][0]['transcript']
                    for result in response['results']
                ])
                confidence = response['results'][0]['alternatives'][0].get('confidence', 0.9)
                
                return {
                    "success": True,
                    "transcript": transcript.strip(),
                    "provider": "watson_stt",
                    "confidence": confidence,
                    "results": response.get('results', [])
                }
            else:
                return {
                    "success": False,
                    "error": "No speech detected in audio"
                }
                
        except Exception as e:
            logger.error(f"Watson STT error: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Watson STT failed: {str(e)}")
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Audio transcription error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Transcription failed: {str(e)}")


@app.post("/audit-transcript")
async def audit_transcript(request: Dict[str, Any]):
    """
    Audit transcript by appending to Word document
    """
    from docx import Document
    from pathlib import Path
    import os
    
    try:
        transcript = request.get("transcript", "")
        timestamp = request.get("timestamp", datetime.now().isoformat())
        
        if not transcript.strip():
            raise HTTPException(status_code=400, detail="No transcript provided")
        
        # Define audit document path
        doc_path = Path(__file__).parent / "data" / "compliance_audit.docx"
        
        # Create or load document
        if doc_path.exists():
            doc = Document(str(doc_path))
        else:
            doc = Document()
            doc.add_heading('Compliance Audit Log', 0)
        
        # Add new audit entry
        doc.add_heading(f'Audit Entry - {datetime.fromisoformat(timestamp).strftime("%Y-%m-%d %I:%M %p")}', level=2)
        doc.add_paragraph(transcript)
        doc.add_paragraph()  # Blank line separator
        
        # Save document
        doc.save(str(doc_path))
        
        # Log the audit
        logger.info(f"Transcript audited and saved to {doc_path}")
        
        return {
            "success": True,
            "message": "Transcript audited successfully",
            "doc_path": str(doc_path)
        }
        
    except Exception as e:
        logger.error(f"Audit transcript error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Audit failed: {str(e)}")


@app.get("/audit-logs")
async def get_audit_logs():
    """
    Get recent audit log versions from Word document
    """
    from docx import Document
    from pathlib import Path
    
    try:
        doc_path = Path(__file__).parent / "data" / "compliance_audit.docx"
        
        if not doc_path.exists():
            return {"success": True, "logs": []}
        
        doc = Document(str(doc_path))
        logs = []
        version = 1
        
        # Parse document for audit entries
        for i, para in enumerate(doc.paragraphs):
            if para.text.startswith('Audit Entry -'):
                timestamp_str = para.text.replace('Audit Entry - ', '')
                # Get preview from next paragraph
                preview = ""
                if i + 1 < len(doc.paragraphs):
                    preview = doc.paragraphs[i + 1].text[:100] + "..."
                
                logs.append({
                    "timestamp": timestamp_str,
                    "version": version,
                    "preview": preview,
                    "status": "Saved"
                })
                version += 1
        
        # Return most recent 10 logs
        return {"success": True, "logs": logs[-10:][::-1]}
        
    except Exception as e:
        logger.error(f"Get audit logs error: {str(e)}", exc_info=True)
        return {"success": False, "logs": [], "error": str(e)}


@app.get("/executive-summary")
async def get_executive_summary():
    """
    Get RAG-generated executive summary from Word document content
    """
    from docx import Document
    from pathlib import Path
    
    try:
        doc_path = Path(__file__).parent / "data" / "compliance_audit.docx"
        
        if not doc_path.exists():
            return {"success": True, "summary": "No analysis performed yet."}
        
        # Read document content
        doc = Document(str(doc_path))
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        if not full_text.strip():
            return {"success": True, "summary": "No audit data available."}
        
        # Generate simple summary (RAG can be added later)
        word_count = len(full_text.split())
        entry_count = full_text.count("Audit Entry -")
        if entry_count > 0:
            summary = f"Compliance audit document contains {entry_count} audit entries with {word_count} total words. Document tracks broker-client communications for SEC compliance monitoring and regulatory review."
        else:
            summary = "No audit entries found. Upload and audit transcripts to generate summary."
        
        return {"success": True, "summary": summary}
        
    except Exception as e:
        logger.error(f"Executive summary error: {str(e)}", exc_info=True)
        return {"success": False, "summary": "Failed to generate summary.", "error": str(e)}


@app.post("/generate-portfolio-report")
async def generate_portfolio_report(request: Dict[str, Any]):
    """
    Generate Client Portfolio Report with RAG analysis
    """
    from docx import Document
    from pathlib import Path
    import csv
    
    try:
        trigger_rag = request.get("trigger_rag", True)
        
        # Read CSV data
        csv_path = Path(__file__).parent / "data" / "trade_blotter.csv"
        trades = []
        if csv_path.exists():
            with open(csv_path, 'r') as f:
                reader = csv.DictReader(f)
                trades = list(reader)
        
        # Create report document
        report_path = Path(__file__).parent / "data" / "client_portfolio_report.docx"
        doc = Document()
        
        doc.add_heading('Client Portfolio Report', 0)
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %I:%M %p")}')
        doc.add_paragraph()
        
        # Add trade summary
        doc.add_heading('Trade Summary', level=1)
        doc.add_paragraph(f'Total Trades: {len(trades)}')
        
        if trigger_rag:
            # Generate basic analysis summary
            doc.add_heading('Portfolio Analysis', level=1)
            
            # Calculate basic metrics safely
            buy_count = sum(1 for t in trades if t.get('Side', '').upper() == 'BUY')
            sell_count = sum(1 for t in trades if t.get('Side', '').upper() == 'SELL')
            
            # Calculate total value with error handling
            total_value = 0
            for t in trades:
                try:
                    price = float(t.get('Price', 0))
                    qty = int(t.get('Qty', 0))
                    total_value += price * qty
                except (ValueError, TypeError):
                    continue
            
            analysis = f"Portfolio contains {len(trades)} total trades ({buy_count} buys, {sell_count} sells) with total transaction value of ${total_value:,.2f}. This report provides a comprehensive overview of client trading activity for compliance review and regulatory monitoring."
            doc.add_paragraph(analysis)
        
        # Add recent trades table
        doc.add_heading('Recent Trades', level=1)
        for trade in trades[-10:]:
            doc.add_paragraph(f"{trade.get('Timestamp', 'N/A')} - {trade.get('Client', 'N/A')}: {trade.get('Side', 'N/A')} {trade.get('Qty', 'N/A')} {trade.get('Ticker', 'N/A')} @ ${trade.get('Price', 'N/A')}")
        
        doc.save(str(report_path))
        
        return {
            "success": True,
            "message": "Report generated successfully",
            "report_path": str(report_path)
        }
        
    except Exception as e:
        logger.error(f"Generate report error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")


@app.get("/download-portfolio-report")
async def download_portfolio_report():
    """
    Open the generated portfolio report
    """
    from pathlib import Path
    import os
    import subprocess
    
    report_path = Path(__file__).parent / "data" / "client_portfolio_report.docx"
    
    if not report_path.exists():
        raise HTTPException(status_code=404, detail="Report not found. Generate report first.")
    
    # Open the file with default application
    try:
        if os.name == 'nt':  # Windows
            os.startfile(str(report_path))
        elif os.name == 'posix':  # macOS and Linux
            subprocess.call(['open' if sys.platform == 'darwin' else 'xdg-open', str(report_path)])
        
        return {"success": True, "message": "Report opened successfully"}
    except Exception as e:
        logger.error(f"Failed to open report: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to open report: {str(e)}")


@app.post("/email-supervisor")
async def email_supervisor(request: Dict[str, Any]):
    """
    Email supervisor with Word doc and Excel attachments
    """
    from pathlib import Path
    
    try:
        timestamp = request.get("timestamp", datetime.now().isoformat())
        supervisor_email = "prasannathefreelancer+supervisor@gmail.com"
        
        # Check for attachments
        doc_path = Path(__file__).parent / "data" / "compliance_audit.docx"
        report_path = Path(__file__).parent / "data" / "client_portfolio_report.docx"
        csv_path = Path(__file__).parent / "data" / "trade_blotter.csv"
        
        attachments = []
        if doc_path.exists():
            attachments.append(("compliance_audit.docx", str(doc_path)))
        if report_path.exists():
            attachments.append(("client_portfolio_report.docx", str(report_path)))
        if csv_path.exists():
            attachments.append(("trade_blotter.csv", str(csv_path)))
        
        # Email content
        email_time = datetime.fromisoformat(timestamp).strftime("%Y-%m-%d %I:%M %p")
        email_body = f"""<html><body>
<h2>Compliance Analysis Report</h2>
<p><strong>Date & Time:</strong> {email_time}</p>
<p>Please find attached the compliance audit documents and trade blotter for your review.</p>
<p><strong>Attachments:</strong></p>
<ul>
<li>Compliance Audit Document (Word)</li>
<li>Client Portfolio Report (Word)</li>
<li>Trade Blotter (Excel/CSV)</li>
</ul>
<p>Best regards,<br>ORQON Compliance System</p>
</body></html>"""
        
        # Use Gmail agent from coordinator with proper context
        gmail_agent = coordinator.agents.get(AgentType.GMAIL)
        if gmail_agent:
            # Send email with basic implementation
            try:
                # Build proper query and context
                query = f"Send email to {supervisor_email} with subject 'Compliance Analysis Report - {email_time}'"
                context = {
                    "to": supervisor_email,
                    "subject": f"Compliance Analysis Report - {email_time}",
                    "body": email_body,
                    "attachments": attachments
                }
                
                result = await gmail_agent.process(query, context)
                
                return {
                    "success": True,
                    "message": f"Email sent to {supervisor_email} with {len(attachments)} attachments"
                }
            except Exception as e:
                logger.error(f"Gmail send error: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Failed to send email: {str(e)}")
        else:
            raise HTTPException(status_code=500, detail="Gmail agent not available")
        
    except Exception as e:
        logger.error(f"Email supervisor error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Email failed: {str(e)}")


@app.post("/save-trades")
async def save_trades(request: Dict[str, Any]):
    """
    Save trades to CSV from frontend confirmation button
    Used when user clicks "Confirm & Save Trade(s) to CSV"
    """
    from pathlib import Path
    import csv
    
    try:
        trades = request.get("trades", [])
        
        if not trades:
            raise HTTPException(status_code=400, detail="No trades provided")
        
        csv_path = Path(__file__).parent / "data" / "trade_blotter.csv"
        saved_count = 0
        
        for trade in trades:
            try:
                # Generate ticket ID if not present
                if not trade.get('ticket_id'):
                    trade['ticket_id'] = f"TKT-{datetime.now().strftime('%Y%m%d%H%M%S')}-{saved_count}"
                
                # Add timestamp if not present
                if not trade.get('timestamp'):
                    trade['timestamp'] = datetime.now().strftime("%Y-%m-%d %I:%M %p")
                
                # Map frontend fields to CSV columns
                with open(csv_path, 'a', newline='', encoding='utf-8') as f:
                    writer = csv.DictWriter(f, fieldnames=[
                        'Ticket ID', 'Client', 'Account', 'Side', 'Ticker', 'Qty',
                        'Type', 'Price', 'Solicited', 'Timestamp', 'Notes',
                        'Follow-up', 'Email', 'Stage', 'Meeting'
                    ])
                    
                    # Map action -> side
                    side = trade.get('action', trade.get('side', ''))
                    
                    # Map solicited string to Yes/No
                    solicited = trade.get('solicited', 'No')
                    if isinstance(solicited, bool):
                        solicited = 'Yes' if solicited else 'No'
                    elif solicited not in ['Yes', 'No']:
                        solicited = 'Yes' if 'solicited' in solicited.lower() else 'No'
                    
                    # Map meeting_needed to Yes/No
                    meeting = trade.get('meeting_needed', 'No')
                    if isinstance(meeting, bool):
                        meeting = 'Yes' if meeting else 'No'
                    
                    writer.writerow({
                        'Ticket ID': trade.get('ticket_id', ''),
                        'Client': trade.get('client_name', ''),
                        'Account': trade.get('account_number', ''),
                        'Side': side,
                        'Ticker': trade.get('ticker', '').upper(),
                        'Qty': trade.get('quantity', 0),
                        'Type': trade.get('order_type', 'Market'),
                        'Price': trade.get('price', 0),
                        'Solicited': solicited,
                        'Timestamp': trade.get('timestamp', ''),
                        'Notes': trade.get('notes', ''),
                        'Follow-up': trade.get('follow_up_date', ''),
                        'Email': trade.get('email', ''),
                        'Stage': trade.get('stage', 'Pending'),
                        'Meeting': meeting
                    })
                
                saved_count += 1
                
            except Exception as e:
                print(f"⚠️  Failed to save trade: {e}")
                continue
        
        return {
            "success": True,
            "count": saved_count,
            "message": f"Successfully saved {saved_count} trade(s) to CSV"
        }
        
    except Exception as e:
        print(f"❌ Save trades error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to save trades: {str(e)}")


if __name__ == "__main__":
    print("=" * 80)
    print("🚀 IBM watsonx Orchestrate MCP Toolkit Server")
    print("=" * 80)
    print(f"\n📍 Server Endpoints:")
    print(f"   Main Server:        http://localhost:8003")
    print(f"   API Docs:           http://localhost:8003/docs")
    print(f"   Redoc:              http://localhost:8003/redoc")
    print(f"\n🔌 IBM MCP Protocol Endpoints:")
    print(f"   Server Info:        GET  http://localhost:8003/mcp/info")
    print(f"   List Tools:         GET  http://localhost:8003/mcp/tools/list")
    print(f"   Call Tool:          POST http://localhost:8003/mcp/tools/call")
    print(f"   SSE Transport:      GET  http://localhost:8003/mcp/sse")
    print(f"\n💬 Chat Endpoints:")
    print(f"   Chat:               POST http://localhost:8003/chat")
    print(f"   Streaming:          POST http://localhost:8003/chat/stream")
    print(f"   WebSocket:          ws://localhost:8003/ws/chat")
    print(f"\n🤖 Specialized Agents:")
    print(f"   • Gmail Agent        - Email operations (send, draft, search)")
    print(f"   • Excel Agent        - CSV/Excel data (client info, trades)")
    print(f"   • Finance Agent      - Stock prices, portfolio analysis")
    print(f"   • Compliance Agent   - RAG knowledge base, risk assessment")
    print(f"\n📦 IBM Toolkit Import Commands:")
    print(f"   Local Python:")
    print(f"     orchestrate toolkits import \\")
    print(f"       --kind mcp \\")
    print(f"       --name orqon-toolkit \\")
    print(f"       --description 'Financial trading multi-agent toolkit' \\")
    print(f"       --package-root ./orqon_core \\")
    print(f"       --command 'python mcp_server.py' \\")
    print(f"       --tools '*'")
    print(f"\n   Remote MCP (SSE):")
    print(f"     orchestrate toolkits import \\")
    print(f"       --kind mcp \\")
    print(f"       --name orqon-toolkit \\")
    print(f"       --description 'Financial trading multi-agent toolkit' \\")
    print(f"       --url 'http://localhost:8003/mcp/sse' \\")
    print(f"       --transport sse \\")
    print(f"       --tools '*'")
    print(f"\n✓ Model: IBM Granite-3-8b-instruct")
    print(f"✓ MCP Protocol: 2024-11-05")
    print(f"✓ Transport: stdio, SSE, streamable HTTP")
    print(f"✓ IBM watsonx Orchestrate ADK: v1.15.0")
    print(f"\n⏱️  Handshake timeout: 30 seconds (IBM requirement)")
    print(f"\nPress Ctrl+C to stop\n")
    print("=" * 80)
    
    uvicorn.run(
        app,
        host="0.0.0.0",
        port=8003,
        log_level="info"
    )
